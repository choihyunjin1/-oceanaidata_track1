"""Build the verified Korean P1/P2/P3 parallel research decision brief."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = (
    ROOT
    / "reports"
    / "parallel_deep_research_execution_20260828_v4"
    / "P1_P2_P3_병렬_딥리서치_결론보고_v4_20260828.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "5F6B7A"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF4EA"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=9.2, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], LIGHT)
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    mark_header_row(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, *, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    # A one-row callout is still represented as a Word table. Mark that row as
    # its header so screen readers do not report an unlabelled table structure.
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(f"{title}  ")
    set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
    run = paragraph.add_run(body)
    set_run_font(run, size=10.2, color=INK)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead:
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=10.5, bold=True, color=INK)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("분당독고다이 | 병렬 딥리서치 실행 보고"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("2026-08-28  |  "), size=9, color=MUTED)
    add_page_field(footer)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    set_run_font(kicker.add_run("TECHNICAL DECISION BRIEF"), size=9.5, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("P1·P2·P3 병렬 딥리서치 결론 보고 v4"), size=23, bold=True, color="000000")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("문제별 구조 탐색, 단일 bounded 실행, 독립 QA, 다음 연구 우선순위"), size=12.5, color=MUTED)
    for label, value in (
        ("팀", "분당독고다이"),
        ("작성일", "2026-08-28 KST"),
        ("판정 범위", "로컬 연구 전용; 공식 test/sample/submission 접근 및 업로드 없음"),
        ("상태", "세 새 후보 모두 NO_GO; 기존 incumbent 유지"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}: "), size=10.5, bold=True, color="000000")
        set_run_font(paragraph.add_run(value), size=10.5, color="000000")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(9)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)


def build(output: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_block(doc)
    add_callout(
        doc,
        "최종 결론",
        "이번 사이클의 새 공식 제출 후보는 없다. P1은 양성 support 부족, P2는 물리 보정의 식별·활성 부족, P3는 ERA5 우위 구간의 범위·신뢰도 부족으로 각각 기각했다. 공식 경로 접근, 후보 CSV 생성, 업로드는 모두 0회다.",
        fill=PALE_RED,
    )
    add_table(
        doc,
        ["문제", "사전등록 구조", "핵심 관측", "결정"],
        [
            ["P1", "동결 83제안 event ranker", "Qualification 양성 1 / event 1", "NO_GO_SUPPORT"],
            ["P2", "Public-only 1-mode heave", "Δ +0.000011°C; 활성 0.0687%", "계열 종료"],
            ["P3", "ERA5 advantage ridge router", "I4 Δ -0.001150m; 활성 1.389%", "I4 NO_GO"],
        ],
        [920, 2600, 3860, 1980],
    )
    add_body(
        doc,
        "공식 종합 점수와 로컬 물리 단위(°C, m)는 서로 다른 척도다. 로컬 delta를 공식 점수 차이로 환산하지 않는다. 이미 노출된 historical surface는 fresh holdout으로 부르지 않았고, 결과 전에 구조·분할·gate·prediction hash를 고정했다.",
        bold_lead="판정 원칙. ",
    )

    doc.add_heading("1. P1 — 83-bank 후처리는 통계적 support에서 종료", level=1)
    add_body(
        doc,
        "기존 83개 proposal 중 73개가 false proposal이므로 얕은 event-level ranker를 검토했다. Neyman-Pearson 분류와 conformal risk control은 recall 또는 위험을 제한하는 선택 규칙의 근거가 되지만, 독립 calibration/qualification에 충분한 사건이 있어야 한다."
    )
    add_table(
        doc,
        ["Support 항목", "관측", "사전 기준", "판정"],
        [
            ["전체 bank", "83 proposals / positive 10", "고정, 재생성 금지", "PASS"],
            ["Train", "33 / positive 5 / event 5", "Positive ≥3", "PASS"],
            ["Calibration", "17 / positive 2 / event 2", "각 ≥2", "PASS"],
            ["Qualification", "14 / positive 1 / event 1", "각 ≥2", "FAIL"],
            ["Normality coverage", "55.42%", "≥80%", "FAIL"],
        ],
        [2540, 2380, 2200, 2240],
    )
    add_callout(
        doc,
        "P1 판정",
        "NO_GO_SUPPORT. 모델 fit 0회, threshold 선택 0회, Q2 truth 0행이다. historical zero-add와 Q2 anchor는 byte-equivalent no-op이다. 결과를 보고 split이나 기준을 완화하지 않았다.",
    )
    add_body(
        doc,
        "83-bank 후처리 계열은 종료한다. 다음 P1은 ranker가 아니라 실제 event group/time split에서 qualification support를 늘리는 proposal generator 구조여야 한다. Synthetic data는 보조 학습에만 쓰고 실제 이벤트 blind 평가를 승격 근거로 둔다.",
        bold_lead="다음 단계. ",
    )

    doc.add_heading("2. P2 — public-only heave는 식별 범위가 너무 좁음", level=1)
    add_body(
        doc,
        "온도 경도에 대한 수직 변위의 1차 접선 보정은 물리적으로 해석 가능하지만, 충분한 public layer span과 안정적인 계절 배경이 필요하다. strongest common OOF인 p2_extrapolated_soft_gate_v2를 결과 전에 comparator로 고정했다."
    )
    add_table(
        doc,
        ["지표", "관측", "승격 기준", "판정"],
        [
            ["RMSE", "0.76836746 → 0.76837870°C", "감소", "FAIL"],
            ["ΔRMSE", "+0.000011241°C", "≤ -0.003°C", "FAIL"],
            ["Bootstrap CI90", "[-0.00000908,+0.00004271]°C", "상한 < 0", "FAIL"],
            ["개선 folds", "1/3", "≥2/3", "FAIL"],
            ["보정 활성", "48/69,850 (0.0687%)", "≥5%", "FAIL"],
            ["Correction max", "0.169967°C", "≤0.20°C", "PASS"],
        ],
        [2540, 2380, 2200, 2240],
    )
    add_callout(
        doc,
        "P2 판정",
        "FAIL_GATE_STOP_NO_CSV_NO_RESEARCH_LOOP. 안전 gate는 통과했지만 성능·CI·fold·활성 gate는 실패했다. 같은 계열의 mode 수, cap, threshold를 결과에 맞춰 바꾸지 않는다.",
    )
    doc.add_heading("P2 QA 해석", level=2)
    add_body(
        doc,
        "최초 독립 QA의 유일한 실패는 candidate-reference-correction에서 1.6879×10⁻¹⁵°C의 부동소수점 뺄셈 차이였다. 예측·지표·bootstrap·gate·hash를 바꾸지 않고 rtol=0, atol=2×10⁻¹⁵로 재검산해 PASS했다."
    )
    add_table(
        doc,
        ["다음 구조", "Bounded 규칙"],
        [
            ["Time-depth latent profile", "Public layer 시간 문맥과 depth basis를 공동 학습"],
            ["Comparator", "0.7683674566°C common OOF를 반드시 사용"],
            ["Preflight", "Endpoint·seasonal support를 결과 전에 검사"],
            ["승격", "Δ≤-0.003°C, CI90 upper<0, ≥2/3 folds"],
        ],
        [2500, 6860],
    )
    add_body(
        doc,
        "내부조석 문헌의 수직 변위·mode 설명력은 물리적 근거일 뿐 P2의 얕고 결측된 층에서 같은 성능을 보장하지 않는다. 이번 실험은 그 전이 한계를 실제 support와 점수로 확인했다.",
        bold_lead="근거의 경계. ",
    )

    doc.add_heading("3. P3 — 우위 신호는 있으나 router support가 불충분", level=1)
    add_body(
        doc,
        "기존 286개 past-only feature, transfer expert, incumbent를 동결했다. station/source/fold/calendar/absolute time을 금지하고, 현재 파고 동역학·주기·에너지·wind proxy·연속 lead·두 예측 차이만 StandardScaler + Ridge(alpha=100)에 입력했다."
    )
    add_table(
        doc,
        ["항목", "관측", "고정 기준", "판정"],
        [
            ["지원 가능 outer", "1/3 folds", "각 I1-I4 support", "부족"],
            ["I4 ΔRMSE", "-0.001150m", "≤ -0.003m", "FAIL"],
            ["I4 CI90", "[-0.003119,0]m", "상한 < 0", "FAIL"],
            ["I4 intervention", "1.389%", "5-50%", "FAIL"],
            ["Outer intervention", "0/1,086 rows", "Gate 통과 시만", "Fallback"],
        ],
        [2740, 2040, 2840, 1740],
    )
    add_callout(
        doc,
        "P3 해석",
        "NO_GO_INNER_I4_GATE. Ridge 1 fit, CatBoost 0 fit이다. I4 gate 실패 뒤 3/3 outer fold가 bit-exact incumbent로 fallback되어 pooled/fold/station/lead delta와 CI90은 모두 0이다.",
        fill=PALE_GREEN,
    )
    add_body(
        doc,
        "ERA5 우위가 일부 구간에서 같은 방향으로 나타났다는 탐색 신호는 보존한다. 다음에는 alpha나 threshold를 완화하지 않고, frozen expert의 forward-sealed historical OOF를 더 넓은 과거 window에 생성해 support를 늘린 뒤 새로운 blind window에서 같은 저용량 router를 한 번 평가한다.",
        bold_lead="다음 단계. ",
    )

    doc.add_heading("4. 공식 제출 판단", level=1)
    add_table(
        doc,
        ["질문", "결론"],
        [
            ["오늘 기회를 쓸 후보가 있는가?", "없음. 세 새 구조 모두 사전 gate 실패"],
            ["작은 로컬 개선을 공식 probe로 볼 수 있는가?", "이번 P3 신호는 magnitude·CI·coverage 동시 실패"],
            ["P2의 안전성만으로 제출 가능한가?", "아님. Strongest comparator 대비 악화"],
            ["P1을 공식으로 확인할 수 있는가?", "아님. 모델 fit 전 support gate 종료"],
        ],
        [3300, 6060],
    )
    add_callout(
        doc,
        "기회비용 원칙",
        "공식 기회는 로컬 delta가 작다는 이유만으로 자동 폐기하지 않되, 사전등록 gate의 magnitude·confidence·support 중 핵심 조건이 동시에 무너진 후보에는 쓰지 않는다. 공식 score와 로컬 단위를 쌍으로 기록하되 배율을 일반화하지 않는다.",
    )

    doc.add_heading("5. 다음 실행 우선순위", level=1)
    add_table(
        doc,
        ["순위", "문제", "다음 구조", "첫 중단 기준"],
        [
            ["1", "P1", "Support-first proposal generator", "Qualification event support <2"],
            ["2", "P2", "Joint time-depth latent profile", "Public endpoint/background support 부족"],
            ["3", "P3", "Forward-sealed expert OOF expansion", "I1-I4 support 미달"],
            ["4", "공통", "Local↔official paired score ledger", "Official pair 없이 환산 금지"],
        ],
        [900, 900, 3860, 3700],
    )
    add_body(
        doc,
        "세 방향은 이번 결과가 보여준 병목을 직접 겨냥한다. 기존 실패 계열의 threshold·cap·blend를 완화하는 재실행은 다음 사이클로 간주하지 않으며 허용하지 않는다.",
        bold_lead="주의. ",
    )

    doc.add_heading("6. 근거 원장", level=1)
    add_table(
        doc,
        ["ID", "1차 출처", "이번 연구에서의 용도"],
        [
            ["P1-1", "Tong et al., NP classification", "제한된 오류율 아래 선택 규칙"],
            ["P1-2", "Angelopoulos et al., CRC / LTT", "위험 통제와 선택-평가 분리"],
            ["P1-3", "Huet et al., Affiliation metrics", "Event-level temporal 평가"],
            ["P2-1", "Geoffroy & Nycander, JGR 2022", "온도 경도와 thermocline displacement"],
            ["P2-2", "Bendinger et al., Ocean Science 2024", "Vertical mode의 적용 경계"],
            ["P3-1", "van der Laan et al., Super Learner", "OOF 기반 결합 원칙"],
            ["P3-2", "Hasson et al., ICML 2023", "시계열 stacking"],
            ["P3-3", "Barber et al., Non-exchangeable CRC", "시간 shift에서의 위험 통제"],
        ],
        [950, 4350, 4060],
    )
    add_body(
        doc,
        "상세 URL과 주장별 근거는 같은 폴더의 report-source.md와 claim-ledger.json에 있다. 문헌의 수치 개선을 해커톤 점수로 직접 이전하지 않았고, 탐색 신호를 incumbent 승리로 과장하지 않았다."
    )

    doc.add_heading("7. 무결성 및 QA — 통합 PASS", level=1)
    add_table(
        doc,
        ["산출물", "SHA-256"],
        [
            ["P1 result.json", "8afca6fc57c7bd98e99478de235533d3c69aa7a994844a71aa466f1c61ec9f4e"],
            ["P2 result.json", "f9626d17833a01f0ae2095eb0eaf2a9c055a16659ac31bc002c589413af52400"],
            ["P3 result.json", "aa9b4931e479c10ee7540a2b688072da433c4353375fa43a673a63946045ec3f"],
            ["P3 sealed predictions", "cf98f7b507008b2150b31cbb992c699ebf88fc04fd80334b8ad7b56c09614cf7"],
        ],
        [2380, 6980],
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
