"""Read-only structural QA for the final Round-E report artifacts."""

from __future__ import annotations

import json
import re
import sys
import zipfile

from docx import Document


docx_path, xlsx_path = sys.argv[1:3]

doc = Document(docx_path)
headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
required_headings = [
    "결론",
    "P1: Router를 셀 단위로 해부한다",
    "P2: 공식 all-row 이차곡선을 층별로 정확히 푼다",
    "P3: 축 오류를 교정하고 hidden Public 곡률을 직접 식별한다",
    "로컬–공식 지표의 신뢰도",
    "동결된 9개 제출 파일",
]

with zipfile.ZipFile(xlsx_path) as archive:
    workbook_xml = archive.read("xl/workbook.xml").decode("utf-8")
    sheet_names = re.findall(r'<(?:\w+:)?sheet[^>]*name="([^"]+)"', workbook_xml)
    formula_nodes = sum(
        len(re.findall(rb"<(?:\w+:)?f(?:\s|>)", archive.read(name)))
        for name in archive.namelist()
        if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
    )
    zip_test = archive.testzip()

required_sheets = {"Summary", "Calibration", "Candidates", "P2_P3_Math", "Protocol_Sources"}
result = {
    "status": "PASS",
    "docx": {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": len(headings),
        "required_headings_present": all(item in headings for item in required_headings),
        "relationships": len(doc.part.rels),
    },
    "xlsx": {
        "zip_test": zip_test,
        "sheets": sheet_names,
        "formula_nodes": formula_nodes,
        "required_sheets_present": required_sheets.issubset(sheet_names),
    },
}
if not result["docx"]["required_headings_present"] or zip_test is not None or not result["xlsx"]["required_sheets_present"]:
    result["status"] = "FAIL"
print(json.dumps(result, ensure_ascii=False, indent=2))
