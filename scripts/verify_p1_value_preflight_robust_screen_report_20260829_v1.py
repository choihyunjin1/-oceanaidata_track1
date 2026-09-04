"""Independent structural and evidence QA for the P1 value-preflight report."""

from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "p1_value_preflight_robust_screen_20260829_v1"
DOCX = REPORT_DIR / "20260829_P1_가치성_사전심사_및_환경강건_저충실도_보고서.docx"
RESULT = ROOT / "artifacts" / "p1_environment_balanced_replay_screen_20260829_v1" / "run" / "result.json"
PREFLIGHT = ROOT / "artifacts" / "p1_environment_balanced_replay_screen_20260829_v1" / "preflight.json"
POSTRUN = ROOT / "artifacts" / "p1_environment_balanced_replay_screen_20260829_v1" / "postrun_gate.json"
QA = REPORT_DIR / "qa.json"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1 << 20), b""):
            digest.update(chunk)
    return digest.hexdigest()


def all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def main() -> None:
    result = json.loads(RESULT.read_text(encoding="utf-8"))
    preflight = json.loads(PREFLIGHT.read_text(encoding="utf-8"))
    postrun = json.loads(POSTRUN.read_text(encoding="utf-8"))
    doc = Document(DOCX)
    text = all_text(doc)
    section = doc.sections[0]

    checks = {
        "preflight_passed_10_of_10": preflight["decision"] == "PASS_TO_LOW_FIDELITY" and preflight["score"] == 10,
        "postrun_stopped": postrun["decision"] == "STOP_BEFORE_FULL_FIDELITY",
        "q3_exact": abs(result["metrics"]["q3"]["delta_f1"] - (-0.0008115153976472333)) < 1e-15,
        "q4_exact": abs(result["metrics"]["q4"]["delta_f1"] - 0.0012710885929277937) < 1e-15,
        "pooled_exact": abs(result["pooled"]["delta_f1"] - 4.261250984927134e-05) < 1e-15,
        "all_added_rows_false_positive": (
            result["metrics"]["q3"]["candidate_added_rows"] + result["metrics"]["q4"]["candidate_added_rows"] == 23
            and result["metrics"]["q3"]["candidate_added_true_rows"] + result["metrics"]["q4"]["candidate_added_true_rows"] == 0
        ),
        "official_boundary": result["official_test_rows_read"] == 0 and not result["submission_created"] and not result["upload_performed"],
        "docx_exists_nonempty": DOCX.exists() and DOCX.stat().st_size > 20_000,
        "letter_page": abs(section.page_width.inches - 8.5) < 0.01 and abs(section.page_height.inches - 11.0) < 0.01,
        "one_inch_margins": all(abs(value.inches - 1.0) < 0.01 for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)),
        "header_footer_distance": abs(section.header_distance.inches - 0.492) < 0.01 and abs(section.footer_distance.inches - 0.492) < 0.01,
        "normal_style": doc.styles["Normal"].font.name == "Calibri" and abs(doc.styles["Normal"].font.size.pt - 11.0) < 0.01,
        "required_headlines_present": all(token in text for token in ("STOP_BEFORE_FULL_FIDELITY", "-0.0008115", "+0.0012711", "+0.0000426", "추가 23행")),
        "claim_limit_present": "retrospective" in text and "full GroupDRO" in text,
        "no_placeholder_text": not any(token in text for token in ("TODO", "TBD", "PLACEHOLDER", "Lorem ipsum")),
        "no_title_bottom_border": True,
    }

    with zipfile.ZipFile(DOCX) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        checks["five_external_sources_linked"] = sum(url in rels_xml for url in (
            "proceedings.mlr.press/v237/awasthi24a.html",
            "jmlr.org/papers/v18/16-558.html",
            "proceedings.mlr.press/v115/wu20a.html",
            "proceedings.mlr.press/v202/he23b.html",
            "arxiv.org/abs/2010.05761",
        )) == 5
        checks["fixed_table_geometry"] = "w:tblLayout w:type=\"fixed\"" in document_xml and "w:tblInd w:w=\"120\"" in document_xml
        first_title_index = document_xml.find("TECHNICAL DECISION MEMO")
        first_heading_index = document_xml.find("1. 이번에 달라진 운영 방식")
        if first_title_index >= 0 and first_heading_index > first_title_index:
            checks["no_title_bottom_border"] = "w:pBdr" not in document_xml[first_title_index:first_heading_index]

    status = "PASS" if all(checks.values()) else "FAIL"
    payload = {
        "schema_version": "p1.value_preflight_report.qa.v1",
        "status": status,
        "checks": checks,
        "hashes": {
            "docx_sha256": sha256(DOCX),
            "result_sha256": sha256(RESULT),
            "preflight_sha256": sha256(PREFLIGHT),
            "postrun_gate_sha256": sha256(POSTRUN),
        },
        "source_files": {
            "docx": str(DOCX.relative_to(ROOT)),
            "result": str(RESULT.relative_to(ROOT)),
            "preflight": str(PREFLIGHT.relative_to(ROOT)),
            "postrun_gate": str(POSTRUN.relative_to(ROOT)),
        },
    }
    QA.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    if status != "PASS":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
