from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUT = (
    REPO
    / "reports"
    / "p2_submit_p1_p3_deep_research_20260828_v1"
    / "P2_공식검증_P1_P3_구조연구_20260828.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "EAF4EA"
RED = "FBECEC"
MUTED = "5F6B7A"


def set_run_font(run, size: float = 11, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, color: str = "000000"):
    p = doc.add_paragraph()
    style_paragraph(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_callout(doc: Document, label: str, text: str, fill: str = CALLOUT, color: str = INK) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, before=6, after=10, line=1.15)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.0)
        set_run_font(p.add_run(value), size=10, bold=True, color=INK)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.0)
            set_run_font(p.add_run(value), size=9.5)
        set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def keep_table_together(table) -> None:
    for row_index, row in enumerate(table.rows):
        keep_next = row_index < len(table.rows) - 1
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = keep_next


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(header, after=0, line=1.0)
    set_run_font(header.add_run("OCEAN AI HACKATHON 2026 | 검증 메모"), size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def build() -> None:
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    style_paragraph(p, before=14, after=4, line=1.0)
    set_run_font(p.add_run("검증 및 구조 연구 메모"), size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    style_paragraph(p, after=4, line=1.0)
    set_run_font(p.add_run("P2 공식 검증 · P1/P3 구조 연구"), size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    style_paragraph(p, after=16, line=1.0)
    set_run_font(p.add_run("분당독고다이 | 2026-08-28 KST | 실험·공식 점수·다음 승격 후보 통합 결론"), size=11, color=MUTED)

    for label, value in (
        ("범위", "P2 공식 1회 검증, P1/P3 신규 실험, 구조적 병목 분석"),
        ("판정", "P2 개선 확정 / P1·P3 로컬 후보 NO_GO / P3 ERA5 과학적 미실행"),
        ("독립 QA", "21/21 PASS · focused pytest 10 PASS · Ruff PASS"),
    ):
        p = doc.add_paragraph()
        style_paragraph(p, after=2, line=1.0)
        set_run_font(p.add_run(f"{label}: "), bold=True, color=INK)
        set_run_font(p.add_run(value), color="000000")

    add_callout(
        doc,
        "핵심 결론",
        "P2 alpha40은 Public RMSE 0.445147°C·27.747847점으로 새 최고가 됐다. P1은 컴퓨팅이 아니라 utility-positive 사건 support가 부족하다. P3 past-only 구조는 장기 lead에서 실패했지만, ERA5 계열은 fit 0회라 아직 모델 실패로 판정할 수 없다.",
        fill=GREEN,
        color=INK,
    )

    add_heading(doc, "1. P2 공식 검증", 1)
    add_table(
        doc,
        ["항목", "직전 최고", "alpha40", "변화"],
        [
            ["Public RMSE (°C)", "0.483661", "0.445147", "-0.038514"],
            ["공식 점수", "27.264587", "27.747847", "+0.483260"],
            ["오늘 잔여", "-", "2 / 3", "2회 보존"],
        ],
        [2700, 2100, 2100, 2460],
    )
    add_body(
        doc,
        "제출 파일은 26,061행·정확한 4열 schema·키 순서·finite/range·PAVA idempotence가 모두 통과했고 SHA-256은 6e28ddb8d78c0969e5104d7efbe28e1762f51e80d759fceb86cdef52baa29b96이다.",
    )
    add_body(
        doc,
        "사전 same-lineage 기하 추정 중심은 0.448627°C였고 실제는 0.445147°C로 0.003480°C 더 좋았다. 반대로 exposed local OOF는 alpha40이 alpha20보다 0.016402°C 나쁘다고 예측했다. 이 축에서는 local surrogate 순위보다 같은 공식 계보의 벡터 기하가 더 잘 운송됐다.",
    )
    add_callout(
        doc,
        "운영 결정",
        "alpha40을 P2 Public incumbent로 보존한다. 남은 두 번을 alpha60/80 자동 연장에 쓰지 않고, fresh holdout 또는 명확히 다른 구조가 생길 때까지 보존한다.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "2. P1: '자원 부족'의 정확한 뜻", 1)
    p1_intro = add_body(
        doc,
        "P1 자원 부족은 GPU·RAM·시간·토큰 부족이 아니다. frozen generator가 만든 proposal 중 학습 가능한 utility-positive 사건의 수와 분포가 부족한 통계 support 문제다.",
    )
    p1_intro.paragraph_format.keep_with_next = True
    p1_support_table = add_table(
        doc,
        ["항목", "관측", "기준", "판정"],
        [
            ["Train utility-positive", "2", "10 이상", "FAIL"],
            ["Calibration utility-positive", "0", "4 이상", "FAIL"],
            ["Hard-negative / positive", "21.5", "5.0 이상", "PASS"],
            ["단일 station×layer 양성 점유", "100%", "70% 이하", "FAIL"],
            ["Verifier fit", "0회", "support 통과 후", "STOP"],
        ],
        [3300, 1800, 1800, 2460],
    )
    keep_table_together(p1_support_table)
    add_body(
        doc,
        "19행 이상 connected proposal은 총 123개였지만 train 45개 중 utility-positive는 2개, calibration 2개 중 utility-positive는 0개였다. epoch를 늘리거나 verifier를 크게 만들어도 이 분모는 늘어나지 않는다.",
    )

    add_heading(doc, "3. P1 NCAD-inspired 실험", 1)
    add_body(
        doc,
        "48시간 causal TCN과 synthetic offset/drift/noise/flatline/contextual outlier exposure로 한 번 검증했다. 이는 NCAD의 정확한 재현이 아니라 아이디어를 현재 데이터에 맞춘 inspired 실험이다.",
    )
    add_table(
        doc,
        ["Surface", "Anchor F1", "Candidate F1", "ΔF1", "추가 TP / FP"],
        [
            ["Inner selection", "0.752104", "0.759129", "+0.007025", "46 / 51"],
            ["Calibration", "0.997912", "0.626474", "-0.371438", "0 / 284"],
            ["Qualification", "0.812060", "0.694456", "-0.117604", "0 / 337"],
        ],
        [2100, 1700, 1700, 1700, 2160],
    )
    add_callout(
        doc,
        "판정: NO_GO_CALIBRATION_SAFETY",
        "합성 anomaly 형태는 selection에서 학습됐지만 station×layer별 실관측으로 운송되지 않았다. 다음 선행조건은 더 큰 모델이 아니라 물리적·station-conditioned positive generator 또는 prospective positive event ledger다.",
        fill=RED,
        color="7A1F1F",
    )

    add_heading(doc, "4. P3: past-only 장기 lead 실험", 1)
    p3_intro = add_body(
        doc,
        "336개 multi-resolution past-only feature와 station별 standardized multi-output ridge를 사용했다. 세 fold 모두 inner alpha 1000을 선택했고, incumbent 보호를 위해 3/6/9h는 고정하고 12/18/24h에 20% residual blend만 적용했다.",
    )
    p3_intro.paragraph_format.keep_with_next = True
    p3_lead_table = add_table(
        doc,
        ["Lead / 범위", "Incumbent RMSE", "Candidate RMSE", "변화"],
        [
            ["Pooled", "0.779949", "0.785851", "+0.005902"],
            ["12h", "0.864363", "0.872553", "+0.008190"],
            ["18h", "0.892958", "0.904090", "+0.011132"],
            ["24h", "0.847421", "0.859850", "+0.012429"],
        ],
        [3000, 2100, 2100, 2160],
    )
    keep_table_together(p3_lead_table)
    add_body(
        doc,
        "Bootstrap 개선확률은 3.62%, 90% CI는 [+0.000438,+0.011481]m였다. 세 station 모두 악화했고 prior TSMixer도 같은 장기 lead에서 악화했다. 단순 용량 부족보다 미래 forcing 정보 부재가 병목이라는 해석이 더 강하다.",
    )
    add_callout(
        doc,
        "판정: TERMINAL_NO_GO",
        "이 past-only linear candidate는 공식 제출 가치가 없다. 그러나 이는 P3 전체 모델 최대치나 exogenous 구조의 실패를 뜻하지 않는다.",
        fill=RED,
        color="7A1F1F",
    )

    add_heading(doc, "5. P3 ERA5: 실패가 아니라 미실행", 1)
    add_table(
        doc,
        ["단계", "상태", "증거"],
        [
            ["Raw download", "PASS", "363 / 363, partial 0"],
            ["Derived / combine", "PASS", "363, 262,917행"],
            ["Preflight", "PASS", "286 features, source quarantine ready"],
            ["Model fit", "미실행", "CatBoost import failure, fit 0회"],
            ["Source/local gate", "미평가", "fit 전 종료"],
        ],
        [2500, 1800, 5060],
    )
    add_body(
        doc,
        "download-only .venv-era5에 competition ML stack이 없어서 ModuleNotFoundError가 발생했다. 모델을 한 번도 fit하지 않았으므로 ERA5 context-transfer가 나쁘다는 과학적 판정은 불가능하다.",
    )
    add_callout(
        doc,
        "다음 승인 후보",
        "새 experiment ID로 frozen 286-feature/source-local split/postprocess/gate를 그대로 유지한다. attempt lock 전에 exact interpreter에서 catboost, sklearn, numpy, pandas, pyarrow import와 버전을 확인한 뒤 실행한다.",
        fill=GREEN,
    )

    add_heading(doc, "6. 문헌과 다음 우선순위", 1)
    add_body(
        doc,
        "NCAD는 contextual outlier exposure로 anomaly problem을 supervised form으로 바꾸는 방향을 제안한다. 이번 P1 결과는 문헌 아이디어 자체가 아니라 local synthetic generator의 현실 운송이 실패했음을 보여준다.",
    )
    add_body(
        doc,
        "DLinear/NLinear은 강한 단순 baseline이지만 이번 P3의 12-24h를 개선하지 못했다. TiDE는 covariate를 포함한 MLP encoder-decoder이고 TimeXer는 exogenous information을 명시적으로 통합한다. 따라서 다음 P3 돌파구는 더 큰 past-only backbone보다 future/exogenous forcing을 정확히 투입하는 것이다.",
    )
    add_table(
        doc,
        ["우선순위", "실행", "승격 전 필수 gate"],
        [
            ["1", "P3 fresh ERA5 frozen-contract attempt", "환경 import/version preflight + 기존 source gate"],
            ["2", "P1 support generation redesign", "분산된 utility-positive support 선확보"],
            ["3", "P2 alpha40 incumbent 보존", "fresh holdout 또는 다른 구조의 실질 개선"],
        ],
        [1300, 3760, 4300],
    )

    add_heading(doc, "7. 재현성 및 제한", 1)
    add_body(doc, "P1/P3 신규 실험은 공식 test/sample/submission을 읽거나 생성하지 않았다.")
    add_body(doc, "P2 점수는 인증된 공식 OCN-02 채점 카드에서 확인했으며 Private metric은 아직 공개되지 않았다.")
    add_body(doc, "P1/P3 로컬 결론은 고정 historical surface에 대한 결론이며 전체 모델 공간의 최대치를 증명하지 않는다.")
    add_body(doc, "독립 QA 21/21, focused pytest 10/10, Ruff가 모두 통과했다.")

    add_heading(doc, "8. 외부 근거", 1)
    for text in (
        "IJCAI 2022 NCAD — https://www.ijcai.org/proceedings/2022/394",
        "DLinear official implementation — https://github.com/honeywell21/DLinear",
        "TiDE primary paper — https://openreview.net/pdf?id=pCbC3aQB5W",
        "TimeXer, NeurIPS 2024 — https://proceedings.neurips.cc/paper_files/paper/2024/file/0113ef4642264adc2e6924a3cbbdf532-Paper-Conference.pdf",
    ):
        add_body(doc, text, color=DARK_BLUE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
