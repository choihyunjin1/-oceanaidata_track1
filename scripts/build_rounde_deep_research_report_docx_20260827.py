"""Build the Korean Round-E structural research report as a polished DOCX.

The Markdown file is the canonical source. This builder only applies a stable
business-brief layout and does not alter analytical content.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "next_day_breakthrough_deep_research_20260827_v1"
SOURCE = REPORT_DIR / "report-source.md"
QA_JSON = REPORT_DIR / "independent_bundle_qa.json"
OUTPUT = REPORT_DIR / "20260827_3x3_구조결함_돌파구_연구보고서_v1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "596574"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F5EC"
GREEN = "276749"
BORDER = "D6DCE5"
WHITE = "FFFFFF"
RED = "9B1C1C"


def set_run_font(run, *, size=None, bold=None, italic=None, color=INK, mono=False):
    latin = "Consolas" if mono else "Arial"
    east = "Malgun Gothic"
    run.font.name = latin
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag_name, attrs in (
        ("tblW", {"w": str(sum(widths_dxa)), "type": "dxa"}),
        ("tblInd", {"w": "120", "type": "dxa"}),
        ("tblLayout", {"type": "fixed"}),
    ):
        node = tbl_pr.find(qn(f"w:{tag_name}"))
        if node is None:
            node = OxmlElement(f"w:{tag_name}")
            tbl_pr.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = borders.find(qn(f"w:{edge}"))
        if line is None:
            line = OxmlElement(f"w:{edge}")
            borders.append(line)
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "5")
        line.set(qn("w:space"), "0")
        line.set(qn("w:color"), BORDER)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    for tag, value in (("color", BLUE), ("u", "single"), ("sz", "19")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        r_pr.append(node)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r_pr.append(fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text, *, size=10.5, color=INK):
    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(8.5, size - 1), color=DARK_BLUE, mono=True)
        else:
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10
    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section):
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    left = header.add_run("해양 해커톤  |  Round E 구조 연구")
    set_run_font(left, size=8.5, bold=True, color=MUTED)
    header.add_run("\t")
    right = header.add_run("분당독고다이  ·  2026-08-27")
    set_run_font(right, size=8.5, color=MUTED)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    footer = section.footer.paragraphs[0]
    left = footer.add_run("내부 사전등록 연구보고서  ·  공식 업로드 전 승인 필요")
    set_run_font(left, size=8.2, color=MUTED)
    footer.add_run("\t")
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    add_page_number(footer)


def add_masthead(doc, qa):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("TECHNICAL RESEARCH BRIEF")
    set_run_font(r, size=9, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("해양 해커톤 다음 날 3×3 제출을 위한\n구조 결함·돌파구 연구")
    set_run_font(r, size=25, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("P1·P2·P3 공식 점수 실험, 로컬–공식 운송성, 2026-08-27 사전등록 후보")
    set_run_font(r, size=12, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [1800, 7560])
    items = [
        ("작성 기준", "2026-08-26 KST"),
        ("팀", "분당독고다이"),
        ("상태", "9개 CSV 동결 · 독립 QA PASS · 공식 업로드 0회"),
    ]
    for i, (label, value) in enumerate(items):
        set_cell_shading(meta.cell(i, 0), PALE_BLUE)
        for para in meta.cell(i, 0).paragraphs:
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(label)
            set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)
        for para in meta.cell(i, 1).paragraphs:
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            set_run_font(run, size=9.5, bold=(label == "상태"), color=GREEN if label == "상태" else INK)
    doc.add_paragraph()

    cards = doc.add_table(rows=1, cols=3)
    set_table_geometry(cards, [3120, 3120, 3120])
    card_data = [
        ("P1", "+0.024163 F1", "G/I/removal 분해"),
        ("P2", "0.535750", "예상 공식 RMSE"),
        ("P3", "0.598987", "C1 guard RMSE"),
    ]
    for idx, (problem, value, note) in enumerate(card_data):
        cell = cards.cell(0, idx)
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(problem)
        set_run_font(r, size=9, bold=True, color=BLUE)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(value)
        set_run_font(r, size=16, bold=True, color=INK)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(note)
        set_run_font(r, size=8.5, color=MUTED)

    doc.add_paragraph()
    qa_box = doc.add_table(rows=1, cols=1)
    set_table_geometry(qa_box, [9360])
    cell = qa_box.cell(0, 0)
    set_cell_shading(cell, PALE_GREEN)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("독립 QA  ")
    set_run_font(r, size=10, bold=True, color=GREEN)
    r = p.add_run(
        f"{qa['status']} · P0/P1/P2 = 0/0/0 · manifest {qa['manifest_sha256'][:12]}…"
    )
    set_run_font(r, size=9.5, color=INK)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("이 문서는 제출 후보와 판정 규칙을 고정한다. 실제 업로드에는 직전의 새 명시 승인이 필요하다.")
    set_run_font(r, size=9, color=MUTED)
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    if cols == 1:
        widths = [9360]
    elif cols == 2:
        widths = [2800, 6560]
    elif cols == 3:
        widths = [2200, 2800, 4360]
    elif cols == 4:
        widths = [1700, 2600, 1500, 3560]
    elif cols == 5:
        widths = [1450, 1900, 1900, 1600, 2510]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx in range(cols):
            cell = table.cell(ridx, cidx)
            if ridx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, row[cidx] if cidx < len(row) else "", size=8.3 if cols >= 5 else 8.8)
            for run in p.runs:
                if ridx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            if cidx > 0 and re.fullmatch(r"[+−\-]?[0-9.,×%\[\] ]+", row[cidx] if cidx < len(row) else ""):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_inline(p, text, size={1: 16, 2: 13, 3: 12}.get(level, 11), color=BLUE if level <= 2 else DARK_BLUE)
    major = (
        "P1:", "P2:", "P3:", "로컬–공식", "동결된 9개", "출처와 내부",
    )
    if level == 2 and any(text.startswith(prefix) for prefix in major):
        p.paragraph_format.page_break_before = True
    return p


def render_markdown(doc, source_text):
    lines = source_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("부제:") or stripped.startswith("작성 기준") or stripped.startswith("팀:") or stripped.startswith("상태:"):
            i += 1
            continue
        heading = re.match(r"^(#{2,3})\s+(.+)$", stripped)
        if heading:
            add_heading(doc, heading.group(2), len(heading.group(1)) - 1)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.50)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            marker = p.add_run(f"{numbered.group(1)}. ")
            set_run_font(marker, size=10.5, bold=True, color=DARK_BLUE)
            add_inline(p, numbered.group(2))
            i += 1
            continue
        block = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("|") or re.match(r"^(-|\d+\.)\s+", nxt):
                break
            block.append(nxt)
            i += 1
        text = " ".join(block)
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        add_inline(p, text)


def build():
    qa = json.loads(QA_JSON.read_text(encoding="utf-8"))
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    doc.core_properties.title = "해양 해커톤 다음 날 3×3 구조 결함·돌파구 연구"
    doc.core_properties.subject = "P1·P2·P3 공식 점수 실험과 로컬–공식 운송성"
    doc.core_properties.author = "분당독고다이"
    doc.core_properties.keywords = "해양 해커톤, P1, P2, P3, local-official calibration, preregistration"
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_masthead(doc, qa)
    render_markdown(doc, source_text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(json.dumps({"status": "PASS", "output": str(OUTPUT)}, ensure_ascii=False))


if __name__ == "__main__":
    build()
