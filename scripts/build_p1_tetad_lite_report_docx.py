"""Build the verified Korean decision memo for the sealed P1 TE-TAD-lite experiment."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_DIR = ROOT / "artifacts" / "p1_tetad_lite_direct_interval_set_v1"
TERMINAL = ARTIFACT_DIR / "terminal_result.json"
SEAL = ROOT / "artifacts" / "p1_tetad_lite_direct_interval_set_v1_execution_seal.json"
OUTPUT = ARTIFACT_DIR / "P1_TE-TAD-lite_딥리서치_실험보고서_20260826.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "596574"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_RED = "FCE8E6"
RED = "9B1C1C"
GREEN = "276749"
BORDER = "D6DCE5"
WHITE = "FFFFFF"


def set_run_font(run, *, size=None, bold=None, italic=None, color=INK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_keep_lines(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, *, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    set_table_borders(table)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, side="bottom", color=BLUE, size=10, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    r_pr.extend((r_fonts, color, underline, size))
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def populate_running_header(header):
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    left = hp.add_run("P1  |  신규 구조 탐색")
    set_run_font(left, size=8.5, bold=True, color=MUTED)
    hp.add_run("\t")
    right = hp.add_run("2026-08-26  ·  검증 완료")
    set_run_font(right, size=8.5, color=MUTED)
    tabs = hp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)


def add_manual_page_header(doc, *, page_break_before=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.space_before = Pt(24)
    paragraph.paragraph_format.space_after = Pt(30)
    paragraph.paragraph_format.line_spacing = 1.0
    left = paragraph.add_run("P1  |  신규 구조 탐색")
    set_run_font(left, size=8.5, bold=True, color=MUTED)
    paragraph.add_run("\t")
    right = paragraph.add_run("2026-08-26  ·  검증 완료")
    set_run_font(right, size=8.5, color=MUTED)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    set_keep_lines(p)
    return p


def add_body(doc, text, *, bold_lead=None, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=color)
        body = p.add_run(text[len(bold_lead) :])
        set_run_font(body, size=11, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=color)
    set_keep_lines(p)
    return p


def add_callout(doc, label, text, *, failure=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.12
    set_paragraph_shading(p, PALE_RED if failure else PALE_BLUE)
    set_paragraph_border(p, side="left", color=RED if failure else BLUE, size=18, space=5)
    lead = p.add_run(label + "  ")
    set_run_font(lead, size=11, bold=True, color=RED if failure else DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body, size=11, color=INK)
    set_keep_lines(p)
    return p


def fill_cell(cell, text, *, bold=False, color=INK, size=9.5, align=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, widths_dxa, *, compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT)
        fill_cell(
            header.cells[idx],
            value,
            bold=True,
            color=DARK_BLUE,
            size=9.0 if compact else 9.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 and len(str(value)) < 28 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(row.cells[idx], value, size=8.8 if compact else 9.4, align=align)
    set_table_geometry(table, widths_dxa)
    before = table._tbl.getprevious()
    if before is not None:
        pass
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    add_manual_page_header(doc, page_break_before=True)


def add_in_section_page_break(doc):
    add_page_break(doc)


def configure_section_geometry(section):
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_document(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    configure_section_geometry(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_page(doc, terminal):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("딥리서치 기반 신규 구조 실험")
    set_run_font(r, size=10.5, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("P1 TE-TAD/TadTR-inspired Lite")
    set_run_font(r, size=23, bold=True, color="000000")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("직접 시간 구간 집합 예측 — 단일 시드 로컬 gate 보고서")
    set_run_font(r, size=13.5, color=MUTED)

    metadata = (
        ("실험 ID", "p1_tetad_lite_direct_interval_set_v1"),
        ("팀", "분당독고다이"),
        ("실행", "2026-08-26 KST · sealed exactly-once"),
        ("최종 판정", terminal["status"]),
        ("범위", "로컬 sanity gate만 수행 · 공식 test/submission 미접근"),
    )
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=10.5, bold=True, color=INK)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5, color=INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_border(rule, side="bottom", color=BLUE, size=12, space=5)

    add_callout(
        doc,
        "결론",
        "현재 구조는 채택하지 않는다. 경계 IoU와 음성창 억제는 통과했지만, 목표 구간 재현율이 0.8235로 사전 기준 0.90에 못 미쳤다. Q2/Q3/Q4는 실행하지 않았고 결과 기반 재튜닝·재시도도 하지 않았다.",
        failure=True,
    )

    add_table(
        doc,
        ("Target recall", "Median IoU", "Negative FP windows", "Final loss"),
        (("0.823529\nFAIL (<0.90)", "0.850564\nPASS", "0\nPASS (≤1)", "0.555521\nfinite"),),
        (2340, 2340, 2340, 2340),
    )

    add_heading(doc, "판단", 1)
    add_body(
        doc,
        "실패는 학습 불능이 아니라 recall 부족이다. 활성화된 query의 위치 추정은 양호했지만, P1의 핵심 병목인 긴 이벤트 누락 회복을 맡기기에는 query 활성화가 부족했다.",
    )
    add_body(
        doc,
        "다음 방향: 같은 모델의 epoch·threshold·loss weight를 즉석 변경하지 않는다. 별도 사전등록으로 target-free change-point proposal bank와 소형 segment rescorer를 검증한다.",
        bold_lead="다음 방향:",
    )


def add_baseline_and_architecture(doc):
    add_heading(doc, "1. 왜 구간 직접 예측을 시험했나", 1)
    add_body(
        doc,
        "동결 Round-B는 전체 frozen OOF 421,032행에서 precision 0.951804, recall 0.792152, F1 0.864670이다. TP/FP/FN/TN은 12,718 / 644 / 3,337 / 404,333이며, FN 3,337개 중 3,330개가 길이 19행 이상의 이벤트에 속했다.",
    )
    add_callout(
        doc,
        "문제 재정의",
        "정밀도보다 재현율이 병목이다. 따라서 개별 행의 확률을 조금 조정하는 대신, 긴 연속 이상을 하나의 시간 구간으로 직접 제안하는 구조가 의미 있는 반증 가능 가설이었다.",
    )
    add_table(
        doc,
        ("기준", "Precision", "Recall", "F1", "FN"),
        (("Frozen Round-B", "0.951804", "0.792152", "0.864670", "3,337"),),
        (2500, 1600, 1600, 1600, 2060),
    )

    add_heading(doc, "2. 구현한 경량 구조", 1)
    add_table(
        doc,
        ("97채널 입력", "Patch", "Encoder", "5 Queries", "비파괴 결합"),
        (("46 수치 + 46 결측 + station/layer/valid", "1,024행 창\n8행 patch", "d=64\n4 heads × 2", "presence +\nstart/end", "Round-B OR\ndetector"),),
        (2340, 1500, 1650, 1650, 2220),
        compact=True,
    )
    add_body(
        doc,
        "각 query는 presence logit과 정규화된 half-open start/end를 낸다. Hungarian matching 후 BCE + 2×endpoint L1 + 2×(1-IoU)를 최적화했다. 19행 미만 proposal은 폐기하고 NMS, gap closing, minimum-run expansion은 쓰지 않았다.",
    )
    add_body(
        doc,
        "명명 경계: 본 모델은 full TE-TAD 재현이 아니라 TE-TAD/TadTR의 query-based set prediction에서 영감을 받은 lite detector다. actual-timeline coordinates, adaptive query selection, decoder self-attention, deformable attention, iterative refinement와 actionness head는 구현하지 않았다.",
        bold_lead="명명 경계:",
    )


def add_protocol_and_result(doc, terminal):
    add_heading(doc, "3. 시간 순서·누수 방지 프로토콜", 1)
    add_body(
        doc,
        "학습은 strict chronological prefix만 사용했고 robust median/IQR도 같은 prefix에서만 적합했다. centered feature lookahead와 창 길이를 함께 고려해 validation 앞 15일을 purge했다. Q2는 threshold 설계 전용, Q3/Q4는 확인용으로 고정했다.",
    )
    add_table(
        doc,
        ("Prefix", "학습 창", "Qualifying events", "Max targets/window", "Purge"),
        (
            ("Q2", "918", "51", "2", "15일"),
            ("Q3", "1,249", "72", "2", "15일"),
            ("Q4", "1,786", "91", "2", "15일"),
        ),
        (1800, 1800, 2100, 2160, 1500),
    )
    add_body(
        doc,
        "확인 fold에서는 blind score NPZ와 SHA-256 receipt를 먼저 기록한 뒤에만 OOF truth를 읽도록 설계했다. 이번 실행은 Q2 이전에 종료되어 이 경로 자체가 호출되지 않았다.",
    )

    add_heading(doc, "4. 실행 결과 — sanity gate", 1)
    sanity = terminal["sanity"]
    rows = (
        ("Finite loss/gradients", "모두 finite", "모두 finite", "PASS"),
        ("Target recall", "≥ 0.90", f"{sanity['target_recall']:.6f} ({sanity['matched_targets']}/{sanity['targets']})", "FAIL"),
        ("Median matched IoU", "≥ 0.75", f"{sanity['median_matched_iou']:.6f}", "PASS"),
        ("Negative FP windows", "≤ 1", str(sanity["negative_window_fp_windows"]), "PASS"),
        ("Final loss", "finite", f"{sanity['final_loss']:.6f}", "PASS"),
    )
    table = add_table(doc, ("검사", "사전 기준", "관측", "판정"), rows, (3300, 1800, 2400, 1860))
    for row in table.rows[1:]:
        verdict = row.cells[3].text.strip()
        if verdict == "FAIL":
            set_cell_shading(row.cells[3], PALE_RED)
            for run in row.cells[3].paragraphs[0].runs:
                set_run_font(run, size=9.4, bold=True, color=RED)
        else:
            for run in row.cells[3].paragraphs[0].runs:
                set_run_font(run, size=9.4, bold=True, color=GREEN)
    add_callout(
        doc,
        terminal["status"],
        "recall gate 하나가 실패해 상태 기계가 TERMINAL_QA로 종결했다. 이는 외부 성능이 낮았다는 검증 결과가 아니라, 본검증에 진입할 최소 구현/학습 능력을 입증하지 못했다는 판정이다.",
        failure=True,
    )


def add_qa_and_interpretation(doc):
    add_heading(doc, "5. 독립 QA", 1)
    add_table(
        doc,
        ("검증 항목", "결과", "근거"),
        (
            ("봉인된 실행 코드", "PASS", "config, runner, model/engine, 3 tests, tinygrad patch의 SHA·bytes 일치"),
            ("동결 입력", "PASS", "feature/label/OOF/anchor 등 8개 입력 SHA·bytes 일치"),
            ("Attempt ordering", "PASS", "execution seal → attempt lock → terminal 순서 정상"),
            ("Gate 독립 재계산", "PASS", "recall 14/17만 실패; IoU·negative FP·finite는 통과"),
            ("Truth firewall", "PASS / 비적용", "Q2 receipt가 없어 OOF truth 평가 미실행"),
            ("배포·제출 산출물", "0개", "terminal JSON 외 blind/threshold/deployment/submission 파일 없음"),
        ),
        (2600, 1700, 5060),
        compact=True,
    )
    add_body(
        doc,
        "독립 감사와 runner 자체 판정 사이에 불일치가 없었다. execution seal SHA-256은 1147d697…4dab3, terminal result SHA-256은 226f508a…db71이다.",
    )

    add_heading(doc, "6. 무엇을 배웠나", 1)
    add_body(
        doc,
        "긍정 신호: 활성화된 query의 median matched IoU 0.8506과 negative-window FP 0은 위치 회귀와 음성창 억제가 작동했음을 보여준다.",
        bold_lead="긍정 신호:",
    )
    add_body(
        doc,
        "결정적 약점: 쉬운 32창 overfit 상황에서도 target 3개를 놓쳤다. 최종 후보는 anchor OR detector이므로 이 구조의 유일한 역할은 FN 회수인데, 바로 그 역할에 대한 최소 recall 기준을 넘지 못했다.",
        bold_lead="결정적 약점:",
    )
    add_body(
        doc,
        "판단: 현재 상태에서 Q2/Q3/Q4를 강행하거나 epoch·threshold를 결과에 맞춰 조정하면 계산비용과 검증 과적합만 늘어난다. 한 번의 사전등록 주기로 이 구조는 종료하는 것이 맞다.",
        bold_lead="판단:",
    )


def add_evidence(doc):
    add_heading(doc, "7. 1차 문헌 근거", 1)
    sources = (
        (
            "TE-TAD · CVPR 2024",
            "장기 temporal interval을 직접 다루며 actual-timeline coordinates와 adaptive query selection을 제안한다. 본 실험의 구간 직접 예측 관점을 지지하지만, 현 lite 구현은 핵심 구성과 decoder self-attention을 생략했다.",
            "https://openaccess.thecvf.com/content/CVPR2024/html/Kim_TE-TAD_Towards_Full_End-to-End_Temporal_Action_Detection_via_Time-Aligned_Coordinate_CVPR_2024_paper.html",
        ),
        (
            "TadTR · 2021",
            "learned query와 set prediction으로 temporal interval을 직접 예측한다. query·Hungarian matching·interval regression의 직접 근거지만, 현 구현은 deformable locality와 refinement/actionness를 생략했다.",
            "https://arxiv.org/abs/2106.10271",
        ),
        (
            "sigmoidF1 · TMLR 2022",
            "F1의 smooth surrogate를 제시해 최종 metric과 학습 목적 정렬의 가능성을 보인다. 이번 실험에는 구현하지 않았으며 interval matching이나 anchor-union 목적과 동일하지 않다.",
            "https://mlanthology.org/tmlr/2022/benedict2022tmlr-sigmoidf1/",
        ),
        (
            "WETAS · ICCV 2021",
            "연속 anomaly points를 가변 길이 segment로 표현하는 관점을 지지한다. 다만 weak supervision과 DTW pseudo-labeling을 쓰므로 이번 dense-label detector와 알고리즘적으로 같지 않다.",
            "https://openaccess.thecvf.com/content/ICCV2021/html/Lee_Weakly_Supervised_Temporal_Anomaly_Segmentation_With_Dynamic_Time_Warping_ICCV_2021_paper.html",
        ),
    )
    for idx, (name, description, url) in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6 if idx > 1 else 0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{idx}. {name}")
        set_run_font(r, size=11.5, bold=True, color=DARK_BLUE)
        add_body(doc, description, after=3)
        link = doc.add_paragraph()
        link.paragraph_format.left_indent = Inches(0.16)
        link.paragraph_format.space_after = Pt(5)
        add_hyperlink(link, "원문 열기", url)

    add_callout(
        doc,
        "근거 해석",
        "네 문헌은 ‘구간 단위 모델링’의 방향을 지지하지만 해양 QC에서의 성능을 직접 보장하지 않는다. 이번 실험은 그 전이 가설을 작은 계산 예산으로 반증한 사례다.",
    )


def add_next_steps(doc):
    add_heading(doc, "8. 다음 연구 주기", 1)
    add_body(
        doc,
        "다음 후보는 target-free change-point proposal bank + 소형 segment rescorer다. 현 query 모델을 재튜닝하는 것이 아니라, 긴 이벤트 후보 생성과 후보 품질 평가를 분리해 recall 병목을 더 직접적으로 겨냥한다.",
    )
    add_table(
        doc,
        ("단계", "고정 설계", "판정 규칙"),
        (
            ("Proposal", "PELT-L1(level/offset), slope/drift, distribution/noise 비용군; min_size=19", "bank를 사전 고정하고 leave-one-bank-out 확인"),
            ("Rescore", "duration, 양 경계 대비, median/MAD, slope, peer residual, anchor support", "정규화 logistic 또는 소형 tree로 제한"),
            ("Output", "frozen Round-B OR approved proposal", "anchor 양성 삭제와 spike/flatline mutation 금지"),
            ("Gate", "inner oracle ceiling → locked outer", "ΔF1≥+0.0255, CI90 low≥+0.012, precision≥0.75, FP/day≤1.05×"),
        ),
        (1650, 4710, 3000),
        compact=True,
    )

    add_heading(doc, "9. 한계", 1)
    add_table(
        doc,
        ("제한", "영향"),
        (
            ("단일 seed·sanity 종료", "일반화 성능이나 seed 안정성을 말할 수 없음"),
            ("Q2/Q3/Q4 미실행", "후보의 역사적 OOF ΔF1·CI·slice 안전성은 미측정"),
            ("Lite 구현", "full TE-TAD/TadTR의 핵심 모듈 효과를 판정한 것이 아님"),
            ("Anchor OR 결합", "기존 FP를 줄일 수 없고 FN 구제에만 기여 가능"),
            ("대상 범위", "19행 이상 offset/drift/noise에만 해당"),
            ("공식 경로 미접근", "리더보드나 제출 성능을 추론할 수 없음"),
        ),
        (2600, 6760),
        compact=True,
    )

    add_heading(doc, "10. 재현성 식별자", 1)
    add_table(
        doc,
        ("항목", "SHA-256"),
        (
            ("Execution seal", "1147d697f94e33a8134402f138a44abadc2d7c6f7ddf53ba77d82ee60ae4dab3"),
            ("Terminal result", "226f508a88ea65a9a3fb415dca33d31cf788189e4081f4c7614eae349643db71"),
            ("Config", "6bebef85c17c406ff4782f5b4a0d630ee746293774033ade6546e0df9a1efb49"),
            ("Runner", "17f8ab15090792cda37618a20691f34fdbeb01ba20c1159373697f80941fe2ad"),
        ),
        (2300, 7060),
        compact=True,
    )
def main():
    terminal = json.loads(TERMINAL.read_text(encoding="utf-8"))
    seal = json.loads(SEAL.read_text(encoding="utf-8"))
    assert terminal["status"] == "NO_GO_IMPLEMENTATION_GATE"
    assert terminal["sanity"]["matched_targets"] == 14
    assert terminal["sanity"]["targets"] == 17
    assert seal["query_preflight"]["result"] == "PASS"

    doc = Document()
    configure_document(doc)
    add_title_page(doc, terminal)
    add_page_break(doc)
    add_baseline_and_architecture(doc)
    add_page_break(doc)
    add_protocol_and_result(doc, terminal)
    add_page_break(doc)
    add_qa_and_interpretation(doc)
    add_page_break(doc)
    add_evidence(doc)
    add_in_section_page_break(doc)
    add_next_steps(doc)

    props = doc.core_properties
    props.title = "P1 TE-TAD/TadTR-inspired Lite 직접 시간 구간 집합 예측"
    props.subject = "딥리서치 기반 신규 구조 탐색 및 사전등록 단일 시드 로컬 실험 보고서"
    props.author = "분당독고다이"
    props.keywords = "P1, temporal interval set prediction, TE-TAD, TadTR, anomaly detection"
    props.comments = "NO_GO_IMPLEMENTATION_GATE; official test/submission paths were not accessed."

    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    main()
