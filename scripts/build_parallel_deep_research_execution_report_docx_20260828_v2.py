from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "parallel_deep_research_execution_20260828_v2"
SOURCE = REPORT_DIR / "report-source.md"
OUTPUT = REPORT_DIR / "20260828_P1_P2_P3_병렬_딥리서치_실행_결론.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "555555"
LIGHT_GRAY = "F2F4F7"
GRID = "B8C2CC"
WHITE = "FFFFFF"
CALIBRI = "Calibri"


def set_font(run, *, size: float = 11, color: str = "000000", bold: bool = False, italic: bool = False) -> None:
    run.font.name = CALIBRI
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), CALIBRI)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), CALIBRI)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "left" if edge == "start" else "right" if edge == "end" else edge
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 11, bold: bool = False) -> None:
    rel_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), CALIBRI)
    fonts.set(qn("w:hAnsi"), CALIBRI)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.extend([fonts, color, underline, sz])
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


TOKEN = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str, *, size: float = 11, base_bold: bool = False, color: str = "000000") -> None:
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("[") and "](http" in part:
            match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", part)
            if match:
                add_hyperlink(paragraph, match.group(1), match.group(2), size=size, bold=base_bold)
                continue
        bold = base_bold
        mono = False
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            bold = True
        elif part.startswith("`") and part.endswith("`"):
            part = part[1:-1]
            mono = True
        run = paragraph.add_run(part)
        set_font(run, size=size, color=color, bold=bold)
        if mono:
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.size = Pt(max(8.5, size - 1))


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=GRAY)


def add_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    existing_abs = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]

    def make_num(abstract_id: int, num_id: int, kind: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), CALIBRI)
        fonts.set(qn("w:hAnsi"), CALIBRI)
        r_pr.append(fonts)
        lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_id))
        num.append(abs_ref)
        numbering.append(num)

    bullet_abs = max(existing_abs, default=0) + 1
    bullet_num = max(existing_num, default=0) + 1
    make_num(bullet_abs, bullet_num, "bullet")
    decimal_abs = bullet_abs + 1
    decimal_num = bullet_num + 1
    make_num(decimal_abs, decimal_num, "decimal")
    return bullet_num, decimal_num


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = CALIBRI
    normal._element.rPr.rFonts.set(qn("w:ascii"), CALIBRI)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CALIBRI)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = CALIBRI
        style._element.rPr.rFonts.set(qn("w:ascii"), CALIBRI)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CALIBRI)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_metadata(document: Document) -> None:
    rows = [
        ("To", "분당독고다이 연구·제출 의사결정"),
        ("From", "P1·P2·P3 병렬 연구 실행"),
        ("Date", "2026-08-28 KST"),
        ("Status", "P2 α40 한 건만 OFFICIAL_PROBE_READY; 업로드 미수행"),
    ]
    for label, value in rows:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f"{label}: ")
        set_font(r, bold=True)
        add_inline(p, value)
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            continue
        rows.append(parts)
    return rows


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    header0 = rows[0][0] if rows else ""
    if cols == 5 and header0 == "문제":
        return [650, 2450, 2550, 1600, 2110]
    if cols == 5:
        return [1450, 1750, 1500, 2350, 2310]
    if cols == 4:
        return [1200, 2800, 2700, 2660]
    if cols == 3:
        return [1600, 3880, 3880]
    return [9360 // cols] * (cols - 1) + [9360 - (9360 // cols) * (cols - 1)]


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, table_widths(rows))
    set_table_borders(table)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for r_idx, values in enumerate(rows):
        tr_pr = table.rows[r_idx]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            shade_cell(cell, LIGHT_GRAY if r_idx == 0 else WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, size=8.7 if len(values) >= 5 else 9.2, base_bold=r_idx == 0)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def build() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("P1·P2·P3 병렬 딥리서치 실행 보고서")
    set_font(r, size=9, color=GRAY, bold=True)
    r = header.add_run("    |    2026-08-28")
    set_font(r, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("분당독고다이  ·  ")
    set_font(r, size=9, color=GRAY)
    add_page_field(footer)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    add_inline(title, "P1·P2·P3 병렬 딥리서치", size=23, base_bold=True, color="000000")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    add_inline(subtitle, "새 승격기준에 따른 실제 실행·독립 QA·공식 probe 판정", size=14, color=GRAY)
    add_metadata(document)

    bullet_num, decimal_num = add_numbering(document)
    lines = source.splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## 결론부터")
    i = start
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = parse_table(table_lines)
            if rows:
                add_table(document, rows)
            continue
        if line.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], size=13, base_bold=True, color=BLUE)
        elif line.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], size=16, base_bold=True, color=BLUE)
        elif re.match(r"^\d+\. ", line):
            p = document.add_paragraph()
            apply_num(p, decimal_num)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = document.add_paragraph()
            apply_num(p, bullet_num)
            add_inline(p, line[2:])
        else:
            p = document.add_paragraph()
            if line.endswith("  "):
                line = line[:-2]
            add_inline(p, line)
        i += 1

    core = document.core_properties
    core.title = "P1·P2·P3 병렬 딥리서치 및 실제 실행 결론"
    core.subject = "새 승격기준, P1/P2/P3 bounded experiments, P2 α40 official probe readiness"
    core.author = "분당독고다이 / Codex"
    core.keywords = "P1, P2, P3, deep research, model validation, official probe"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
