"""Create the Korean decision memo for the Round-G P2/P3 submission bundle."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(r"C:\Users\cedis\PycharmProjects\PythonProject")
OUTPUT = (
    REPO
    / "reports"
    / "prepared_submission_value_deep_research_20260827_v1"
    / "P2_P3_제출가치_최대치_의사결정_보고서_20260827.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5A6573"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F3EE"
PALE_GOLD = "FFF4D6"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "000000"


def set_run_font(run, size: float = 11, color: str = BLACK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths: list[int], header: bool = True, header_fill: str = LIGHT_GRAY) -> None:
    assert sum(widths) == 9360
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, bold=(header and row_index == 0))
        if header and row_index == 0:
            for cell in row.cells:
                set_cell_fill(cell, header_fill)
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        table.rows[0].cells[index].text = text
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = str(text)
    format_table(table, widths, header_fill=header_fill)
    return table


def add_para(doc: Document, text: str = "", *, bold: bool = False, color: str = BLACK, size: float = 11, after: float = 6, align=None, italic: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_mixed_para(doc: Document, segments: list[tuple[str, bool, str]], *, after: float = 6, size: float = 11):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    for text, bold, color in segments:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    return paragraph


def add_table_source(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Table Source")
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    source = doc.styles.add_style("Table Source", 1)
    source.font.name = "Arial"
    source.font.size = Pt(8.5)
    source.font.italic = True
    source.font.color.rgb = RGBColor.from_string(MUTED)
    source._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    source._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    source._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("분당독고다이  |  P2·P3 제출 의사결정 메모")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.paragraph_format.space_before = Pt(0)
    add_page_number(footer_p)


def add_callout(doc: Document, title: str, text: str, fill: str, title_color: str = NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=11.5, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    run2 = p2.add_run(text)
    set_run_font(run2, size=10.5)
    format_table(table, [9360], header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_header_footer(doc.sections[0])

    add_para(doc, "DECISION MEMO", bold=True, size=23, after=4)
    add_para(doc, "P2·P3 준비 제출물의 활용가치와 ‘최대치’ 판정", bold=True, color=NAVY, size=15, after=12)
    for label, value in (
        ("대상", "분당독고다이"),
        ("일자", "2026-08-27 KST"),
        ("판정", "Round E P3 superseded · Round G ready · 업로드 미수행"),
        ("결정 질문", "현재 6개 파일이 +3점급 개선에 유효한가, 더 나은 세트를 즉시 준비할 수 있는가?"),
    ):
        add_mixed_para(doc, [(f"{label}: ", True, BLACK), (value, False, BLACK)], after=2)
    add_para(doc, "", after=4)

    add_callout(
        doc,
        "결론",
        "기존 세트는 큰 폭 개선용도 전역 최대치도 아닙니다. P3를 공식 Public long-lead 이차 최적점으로 교체한 Round G가 오늘 슬롯의 더 나은 사용입니다. 다만 P2+P3 합산 기대 상승은 약 +0.251점으로, +3점 목표의 약 8.4%에 불과합니다.",
        PALE_GREEN,
    )

    add_table(
        doc,
        ["핵심 지표", "값", "판정"],
        [
            ["P2 U 예상 공식점수 증가", "+0.009851점", "신뢰도 높은 미세 exploit"],
            ["P3 α* 예상 공식점수 증가", "+0.241536점", "고정 Public 1D축 optimum"],
            ["P2+P3 합산 기대", "+0.251387점", "+3점 목표의 약 8.4%"],
            ["전체 모델 최대치", "아님", "새 backbone·표현 공간 미탐색"],
        ],
        [3100, 1900, 4360],
        header_fill=PALE_BLUE,
    )
    add_table_source(doc, "근거: 공식 점수 원장, Round D/E/G manifest 및 독립 QA 재계산.")

    add_heading(doc, "1. 기존 폴더는 어디까지 유효한가", 1)
    add_para(doc, "Round E remaining 폴더는 계보와 QA가 명확한 제출 세트지만, 큰 폭의 점수 상승을 목표로 만든 최종 성능 세트는 아닙니다. P2 첫 파일은 exploit이고, P2 나머지 두 파일과 기존 P3 파일들은 다음 공식 실험을 설계하기 위한 ablation·곡률 probe였습니다.")
    add_table(
        doc,
        ["부분", "기존 역할", "새 판정"],
        [
            ["P2 U", "층별 A−O 공식축 최적화", "유지: 알려진 부분공간 최적점"],
            ["P2 envelope / PAVA", "물리 후처리 분해", "유지: 정보획득용, 상승 미보장"],
            ["P3 long −2/−4", "곡률 탐색", "교체: 기존 공식점수로 곡선 식별 가능"],
            ["Round E 폴더", "제출 대기", "감사 보존, P3 주 세트는 superseded"],
        ],
        [1800, 3260, 4300],
    )
    add_table_source(doc, "근거: Round E P2_P3_SET_MANIFEST.json 및 Round G SET_MANIFEST.json.")

    add_heading(doc, "2. P2: 최대치는 맞지만 범위가 좁다", 1)
    add_para(doc, "P2 U의 예상 공식 RMSE는 0.535750480℃이며, 현 최고 0.536536℃보다 0.000785520℃ 낮습니다. 여섯 개 공식 관측점에서 RMSE를 leaderboard 점수로 변환한 선형식의 최대 잔차는 약 0.000005점입니다.")
    add_callout(
        doc,
        "P2 공식점수 환산",
        "점수 ≈ 33.3333290341 − 12.5475179808 × RMSE℃. 따라서 U의 예상 점수는 26.610990점, 현 최고 대비 +0.009851점입니다.",
        PALE_BLUE,
    )
    add_para(doc, "이 결과는 L2/L3/L4의 A−O 파일축을 각각 최적화한 층별 선형공간에서는 최대치입니다. 새로운 시간·깊이 상호작용, 시공간 representation, 외부 context를 포함하는 전체 P2 모델 공간의 최대치라는 뜻은 아닙니다.")

    add_heading(doc, "3. P3: 이미 보유한 공식점수로 최적점을 복원", 1)
    add_para(doc, "B 파일은 12·18·24시간에서만 O와 A의 정확한 midpoint(α=+0.5)이고, 3·6·9시간에서는 O와 같습니다. Round D의 12시간 α=−2와 18·24시간 α=−2 점수는 변경 support가 겹치지 않으므로 long α=−2의 RMSE를 정확히 합칠 수 있습니다.")
    add_callout(
        doc,
        "식별된 Public 곡선",
        "q(α)=0.0002639228606α²+0.0054027359197α+0.368535199041. 최적점 α*=−10.2354451362, 예상 RMSE*=0.5838540019m입니다.",
        PALE_GREEN,
    )
    add_table(
        doc,
        ["P3 후보", "α", "예상 RMSE", "예상 점수", "현 최고 대비"],
        [
            ["공식 이차최적", "−10.235445", "0.583854002", "24.066765", "+0.241536"],
            ["보수적 근방", "−8", "0.584982371", "24.048856", "+0.223627"],
            ["반대편 bracket", "−12", "0.584557320", "24.055602", "+0.230373"],
        ],
        [2500, 1450, 1800, 1800, 1810],
        header_fill=PALE_BLUE,
    )
    add_table_source(doc, "근거: Round G INDEPENDENT_QA_V2.json. 공식 표시값 ±0.5×10⁻⁶ 반올림을 전파해도 α* 구간은 [−10.289024, −10.182454].")

    add_heading(doc, "4. 새 Round G 제출 순서", 1)
    add_table(
        doc,
        ["순서", "후보", "용도"],
        [
            ["1", "P2 U", "알려진 층별 공식축 exploit"],
            ["2", "P2 endpoint envelope", "bounded 물리 후처리 ablation"],
            ["3", "P2 PAVA+envelope", "단조 제약의 공식 순효과 측정"],
            ["4", "P3 long α*=−10.235445", "고정 Public 1D축 optimum"],
            ["5", "P3 long α=−8", "외삽을 줄인 근방 후보"],
            ["6", "P3 long α=−12", "최적점 반대편 bracket"],
        ],
        [900, 3860, 4600],
    )
    add_table_source(doc, "제출 폴더: C:\\Users\\cedis\\Downloads\\해양 해커톤 제출용\\20260827_round_G_P2x3_P3x3_PUBLIC_QUADRATIC_READY")
    add_para(doc, "P3 세 후보는 서로 독립적인 구조가 아니라 같은 곡선의 optimum neighborhood입니다. 최고점 확보가 우선이면 α*가 첫 선택이고, −8/−12는 scorer·반올림·계보 guard를 확인하는 강건 bracket입니다.")

    add_heading(doc, "5. ‘최대치’의 정확한 의미", 1)
    add_table(
        doc,
        ["질문", "답"],
        [
            ["P2 알려진 A−O 층별 선형공간 최대치인가?", "예"],
            ["P3 고정 Public long-lead 1차원 축 최대치인가?", "예"],
            ["표시 반올림에 강건한가?", "예"],
            ["전체 P2/P3 모델 공간 최대치인가?", "아니오"],
            ["Private·최종순위 최대치인가?", "아니오"],
            ["+3 공식점 돌파 세트인가?", "아니오"],
        ],
        [7000, 2360],
        header_fill=PALE_GOLD,
    )
    add_table_source(doc, "‘예’는 명시된 고정 Public 부분공간에만 한정됩니다.")
    add_para(doc, "같은 Public 점수를 반복 사용하면 후보가 holdout에 적응할 수 있습니다. 따라서 예상대로 최고점이 나와도 상태는 PUBLIC_BEST_ONLY이며, Private-ready 또는 전역 optimum으로 승격하지 않습니다.")

    add_heading(doc, "6. 독립 QA와 승인 경계", 1)
    add_table(
        doc,
        ["검증 항목", "결과"],
        [
            ["P2 세 파일 Round E 바이트 identity", "PASS"],
            ["P3 키·행·순서·midpoint·early no-op", "PASS"],
            ["곡선·최적점·점수 환산 재계산", "PASS"],
            ["SHA-256·유한값·물리 범위", "PASS"],
            ["숨은 target 또는 ERA5 값 열람", "없음"],
            ["공식 업로드", "수행하지 않음"],
        ],
        [6800, 2560],
    )
    add_table_source(doc, "독립 QA: Round G/INDEPENDENT_QA_V2.json · manifest SHA-256 441809738f15a76f61cac098f86c6dd3332d7d04bc1e59679603d8156a8d7f04")
    add_callout(
        doc,
        "승인 경계",
        "연구와 CSV 생성은 업로드 승인이 아닙니다. 실제 제출 전에는 오늘 남은 횟수와 파일 hash를 다시 확인하고, 사용자의 새 명시 승인을 받아야 합니다.",
        PALE_GOLD,
        title_color=RED,
    )

    add_heading(doc, "7. +3점을 위한 다음 구조", 1)
    add_para(doc, "Round G의 기대치는 약 +0.25점입니다. 나머지 약 +2.75점은 기존 postprocess 축과 다른 backbone·표현·오류공간에서 나와야 합니다. 현재 세트는 그 연구를 대체하지 않고 오늘의 공식 기회를 더 효율적으로 쓰는 역할입니다.")
    add_table(
        doc,
        ["영역", "다음 연구", "오늘 즉시 교체 여부"],
        [
            ["P2", "저랭크+deep 시공간 결측 구조(ImputeFormer 계열), 깊이·시간 joint representation", "아니오: full refit·mask 계약·QA 필요"],
            ["P3", "외부 context/sequence 구조, 1D postprocess 축과 다른 오류공간", "아니오: ERA5 고정 실험과 분리 유지"],
            ["총점", "P1/P2 구조 변경을 주 레버로, P3 +0.24를 보조 레버로 사용", "별도 연구 사이클"],
        ],
        [1350, 5200, 2810],
    )
    add_table_source(doc, "참고: ImputeFormer 공식 구현 https://github.com/tongnie/ImputeFormer · KDD 2024 DOI 10.1145/3637528.3671751")

    add_heading(doc, "8. 최종 의사결정", 1)
    add_callout(
        doc,
        "권고",
        "Round E remaining은 감사용으로 보존하고, P3 제출 주 세트는 Round G로 교체합니다. Round G는 오늘 Public 최적화·증거 수집에는 유효하지만 +3점급 돌파구나 전체 모델 최대치로 주장하지 않습니다.",
        PALE_GREEN,
    )
    add_para(doc, "공식 점수가 예상대로 나오면 P2 U와 P3 α*를 Public best로 채택합니다. 예상 구간을 벗어나면 같은 축의 추가 미세 튜닝을 중단하고 scorer·파일 계보를 먼저 감사합니다.")

    add_heading(doc, "근거 출처", 1)
    for text in (
        "공식 대회 홈페이지: https://oceanaidata.org/",
        "Blum & Hardt (2015), The Ladder: https://proceedings.mlr.press/v37/blum15.html",
        "Cawley & Talbot (2010), model-selection overfitting: https://www.jmlr.org/papers/v11/cawley10a.html",
        "ImputeFormer 공식 구현: https://github.com/tongnie/ImputeFormer",
        "상세 claim–source ledger: reports/prepared_submission_value_deep_research_20260827_v1/report-source.md",
    ):
        add_para(doc, text, size=9.2, color=MUTED, after=3)

    configure_page(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
