"""Build the P1 value-preflight and low-fidelity screening report DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "p1_value_preflight_robust_screen_20260829_v1"
OUTPUT = REPORT_DIR / "20260829_P1_가치성_사전심사_및_환경강건_저충실도_보고서.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "666666"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
POSITIVE = "1F3A5F"
CAUTION = "7A5A00"
RISK = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, size: float = 11, bold: bool | None = None,
                 color: str = "000000", italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80,
                     start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        elem = tc_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError("table widths must sum to 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def format_paragraph(paragraph, before: float = 0, after: float = 6,
                     line_spacing: float = 1.10,
                     keep_with_next: bool | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next


def add_text(doc, text: str, bold_prefix: str | None = None,
             color: str = "000000", after: float = 6) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True, color=color)
        set_run_font(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        set_run_font(p.add_run(text), color=color)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    format_paragraph(p, after=8, line_spacing=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_run_font(p.add_run(text))


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    format_paragraph(p, after=8, line_spacing=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_run_font(p.add_run(text))


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level],
                 bold=True, color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])


def add_callout(doc, title: str, body: str, fill: str = BLUE_GRAY,
                title_color: str = INK) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    format_paragraph(p, after=3)
    set_run_font(p.add_run(title + "\n"), size=12, bold=True, color=title_color)
    set_run_font(p.add_run(body), size=10.5, color=INK)
    spacer = doc.add_paragraph()
    format_paragraph(spacer, after=2)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.extend([r_pr, text_elem])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(hp, after=0)
    set_run_font(hp.add_run("P1 연구 운영 보고서  |  가치성 사전심사"), size=9, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, before=4, after=4)
    set_run_font(p.add_run("TECHNICAL DECISION MEMO"), size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    format_paragraph(p, after=5)
    set_run_font(p.add_run("P1 가치성 사전심사 및\n환경강건 저충실도 검증"), size=23, bold=True, color=INK)
    p = doc.add_paragraph()
    format_paragraph(p, after=14)
    set_run_font(p.add_run("가치 없는 실험을 계산 전에 또는 수 분 안에 중단하는 단계형 승격 체계"), size=12.5, color=MUTED)

    for label, value in (
        ("일자", "2026-08-29"),
        ("범위", "P1 retrospective local shadow only"),
        ("판정", "STOP_BEFORE_FULL_FIDELITY"),
        ("운영 경계", "official test 0 · submission 0 · upload 0"),
    ):
        p = doc.add_paragraph()
        format_paragraph(p, after=2)
        set_run_font(p.add_run(f"{label}: "), bold=True, color=INK)
        set_run_font(p.add_run(value), color=INK)

    doc.add_paragraph()


def add_results_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [1200, 1600, 1600, 1560, 1700, 1700]
    set_table_geometry(table, widths)
    headers = ["구간", "baseline F1", "candidate F1", "delta F1", "추가 행 / true", "제거 행 / true"]
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        format_paragraph(p, after=0)
        set_run_font(p.add_run(text), size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    rows = (
        ("Q3", "0.9067084", "0.9058969", "-0.0008115", "4 / 0", "13 / 9"),
        ("Q4", "0.8872198", "0.8884909", "+0.0012711", "19 / 0", "31 / 0"),
        ("pooled", "0.8987428", "0.8987854", "+0.0000426", "23 / 0", "44 / 9"),
    )
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, values, strict=True)):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if index else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, after=0)
            color = RISK if text.startswith("-") else (POSITIVE if text.startswith("+") else "000000")
            set_run_font(p.add_run(text), size=9.5, color=color, bold=(index == 3))


def add_gate_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [700, 1700, 3660, 3300]
    set_table_geometry(table, widths)
    headers = ["단계", "비용", "통과 조건", "실패 시"]
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        format_paragraph(p, after=0)
        set_run_font(p.add_run(text), size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    rows = (
        ("0", "초", "계약·신규성 10개 검사 전부 통과", "폐기 또는 계약 보정"),
        ("1", "수 분", "Q3·Q4·pooled 모두 양수 + 집중도 통과", "즉시 STOP"),
        ("2", "수십 분", "30 epoch에서도 방향 유지 + 추가 precision > 0", "full 금지"),
        ("3", "시간", "3-seed 평균·worst slice·pooled 비열화 없음", "공식 후보 금지"),
        ("4", "희소 기회", "근거와 정보가치 충분 + 사용자 승인", "괴리 원장만 갱신"),
    )
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, values, strict=True)):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            format_paragraph(p, after=0)
            set_run_font(p.add_run(text), size=9.3)


def add_source(doc: Document, number: int, label: str, url: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=4)
    set_run_font(p.add_run(f"{number}. {label} — "), size=9.5, bold=True, color=INK)
    add_hyperlink(p, url, url)


def build() -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "결론",
        "정적 가치성 심사와 10-epoch 저충실도 gate를 운영 기준으로 채택한다. 이번 환경균형 replay는 pooled가 +0.000043이었지만 Q3가 하락하고 추가 23행이 모두 오탐이어서 본학습 전에 중단한다.",
        fill=BLUE_GRAY,
    )

    add_heading(doc, "1. 이번에 달라진 운영 방식", 1)
    add_text(doc, "가치 없는 실험을 완전히 예측할 수는 없다. 대신 계산 비용과 증거 강도를 분리해, 싼 단계에서 실패한 후보가 비싼 단계로 넘어가지 못하게 한다.")
    flow = doc.add_table(rows=2, cols=4)
    flow.style = "Table Grid"
    set_table_geometry(flow, [2340, 2340, 2340, 2340])
    for idx, label in enumerate(("정적 preflight", "10-epoch rung", "30-epoch rung", "3-seed full")):
        set_cell_shading(flow.cell(0, idx), BLUE_GRAY if idx < 2 else LIGHT_GRAY)
        p = flow.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, after=0)
        set_run_font(p.add_run(label), size=10, bold=True, color=INK)
    for idx, label in enumerate(("PASS", "STOP", "미실행", "미실행")):
        p = flow.cell(1, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, after=0)
        set_run_font(p.add_run(label), size=10, bold=True, color=POSITIVE if idx == 0 else RISK if idx == 1 else MUTED)

    add_heading(doc, "2. 정적 가치성 preflight", 1)
    add_text(doc, "다음 항목은 GPU를 쓰기 전에 자동 검사한다.")
    for text in (
        "공식 목적과 같은 binary row F1, 확인 시기 2개 이상, G/I/S station slice를 요구한다.",
        "저충실도 단계·명시적 중단 조건·근거 문헌·반증 근거를 요구한다.",
        "공식 test label, sample/submission, CSV 생성·업로드를 계약에서 금지한다.",
        "12개 닫힌 실험 축과 intervention layer·mechanism 태그가 중복되면 자동 탈락한다.",
    ):
        add_bullet(doc, text)
    add_callout(doc, "정적 판정", "10/10 · PASS_TO_LOW_FIDELITY · 단, Q3/Q4는 retrospective window이므로 fresh claim은 금지", fill=LIGHT_GRAY)

    add_heading(doc, "3. 실제 저충실도 결과", 1)
    add_results_table(doc)
    p = doc.add_paragraph()
    format_paragraph(p, before=4, after=4)
    set_run_font(p.add_run("표 근거: "), size=9, bold=True, color=MUTED)
    set_run_font(p.add_run("artifacts/p1_environment_balanced_replay_screen_20260829_v1/run/result.json"), size=9, color=MUTED)
    add_text(doc, "pooled 미세 양수만 보면 승격처럼 보이지만, Q3에서는 제거한 13행 중 9행이 실제 양성이었고 Q4에서는 제거한 31행이 모두 오탐이었다. 같은 조정이 시기별로 반대 역할을 했으므로 안정적인 환경강건 신호가 아니다.")
    add_callout(doc, "기계 판정", "STOP_BEFORE_FULL_FIDELITY — full GroupDRO, 3-seed 장기학습, 공식 제출 후보 생성을 실행하지 않는다.", fill="FDECEC", title_color=RISK)

    add_heading(doc, "4. 새 승격 기준", 1)
    add_gate_table(doc)
    add_text(doc, "소수점 개선 자체를 무시하지는 않는다. 작은 개선일수록 시기 일관성, 추가 행 precision, station 집중도, seed 안정성을 함께 요구한다. 로컬 F1과 공식 점수는 단위가 다르므로 직접 환산하지 않는다.")

    add_heading(doc, "5. 문헌 근거와 반증", 1)
    add_text(doc, "GroupDRO는 사전에 정의된 그룹의 최악 성능을 낮추려는 목적을 가진다. 환경 정의가 타당하다는 근거이지 이번 proxy의 성공을 보장하지는 않는다.")
    add_text(doc, "Hyperband와 multi-fidelity 최적화는 적은 epoch·데이터 같은 싼 평가로 약한 후보를 조기 중단하고 자원을 유망 후보에 집중한다. 이번 10→30→full 구조의 직접적 근거다.")
    add_text(doc, "반대로 IRM은 가정이나 환경 구조가 맞지 않으면 ERM을 이기지 못할 수 있다. 불변성이라는 이름만으로 후보를 승격하지 않는 이유다.")
    add_text(doc, "RAINCOAT의 time-frequency alignment는 다른 구조적 축이지만, P1에서는 target covariate·transductive 허용성과 leakage 경계를 먼저 잠가야 한다.")

    add_heading(doc, "6. 한계", 1)
    for text in (
        "Q3/Q4는 이미 노출된 retrospective window이므로 fresh generalization 주장이 아니다.",
        "10 epoch와 full ranking의 상관이 낮으면 late bloomer를 거짓 탈락시킬 수 있다.",
        "station별 추가 행 집계는 schema 전환 전에 실행되어 unavailable이며, 낙관적으로 보간하지 않고 fail-closed 처리했다.",
        "단일 seed는 분산을 추정하지 못하며, 이번 실험은 full GroupDRO가 아니라 균형 replay proxy다.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "7. 다음 행동", 1)
    for text in (
        "환경균형 replay family는 닫고 registry에 유지한다.",
        "다음 가설은 time-frequency target representation alignment 또는 명시적 worst-environment loss 중 하나로 한정한다.",
        "새 가설은 정적 preflight와 10→30 epoch 두 rung을 통과한 경우에만 3-seed full로 승격한다.",
        "공식 probe는 로컬 증거와 정보가치가 충분하고 사용자가 정확한 파일을 승인한 경우에만 수행한다.",
    ):
        add_number(doc, text)
    add_callout(doc, "이번 사이클에서 승격된 것", "새 모델이 아니라 저가치 실험을 빠르게 중단하는 운영 장치다.", fill=BLUE_GRAY)

    add_heading(doc, "8. 출처", 1)
    add_source(doc, 1, "Awasthi et al. — Group DRO", "https://proceedings.mlr.press/v237/awasthi24a.html")
    add_source(doc, 2, "Li et al. — Hyperband", "https://www.jmlr.org/papers/v18/16-558.html")
    add_source(doc, 3, "Wu et al. — Multi-fidelity BO", "https://proceedings.mlr.press/v115/wu20a.html")
    add_source(doc, 4, "He et al. — RAINCOAT", "https://proceedings.mlr.press/v202/he23b.html")
    add_source(doc, 5, "Rosenfeld et al. — IRM counterevidence", "https://arxiv.org/abs/2010.05761")

    add_heading(doc, "9. 재현 파일", 1)
    reproduce_files = (
        "configs/experiments/p1_environment_balanced_replay_screen_20260829_v1.json",
        "configs/experiments/p1_experiment_value_registry_20260829_v1.json",
        "artifacts/p1_environment_balanced_replay_screen_20260829_v1/preflight.json",
        "artifacts/p1_environment_balanced_replay_screen_20260829_v1/run/result.json",
        "artifacts/p1_environment_balanced_replay_screen_20260829_v1/postrun_gate.json",
    )
    p = doc.add_paragraph()
    format_paragraph(p, after=0, line_spacing=1.0)
    for index, text in enumerate(reproduce_files):
        set_run_font(p.add_run(text + ("\n" if index < len(reproduce_files) - 1 else "")), size=8.5, color=MUTED)

    doc.core_properties.title = "P1 가치성 사전심사 및 환경강건 저충실도 검증"
    doc.core_properties.subject = "P1 experiment value preflight and low-fidelity gate"
    doc.core_properties.author = "분당독고다이"
    doc.core_properties.keywords = "P1, preflight, multi-fidelity, GroupDRO, experiment gate"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
