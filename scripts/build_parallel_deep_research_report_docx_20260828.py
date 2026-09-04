"""Build the 2026-08-28 parallel P1/P2/P3 Deep Research decision memo.

The Markdown file is the canonical analytical source.  This builder applies
the documents-skill `decision_memo` preset (standard_business_brief alias)
and the `memo_masthead` first-page pattern.
"""

from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "parallel_deep_research_20260828_v1"
SOURCE = REPORT_DIR / "report-source.md"
DECISIONS = REPORT_DIR / "decision_matrix.json"
OUTPUT = REPORT_DIR / "20260828_P1_P2_P3_병렬_딥리서치_의사결정보고서_v1.docx"
AUDIT = REPORT_DIR / "docx_preset_audit.json"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "596574"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F5EC"
PALE_GOLD = "FFF8E1"
GREEN = "276749"
GOLD = "7A5A00"
BORDER = "D6DCE5"
WHITE = "FFFFFF"

INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))")


def set_run_font(run, *, size=11, bold=None, italic=None, color=INK, mono=False):
    latin = "Consolas" if mono else "Arial"
    east = "Malgun Gothic"
    run.font.name = latin
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag_name, attrs in (
        ("tblW", {"w": str(sum(widths_dxa)), "type": "dxa"}),
        ("tblInd", {"w": "120", "type": "dxa"}),
        ("tblLayout", {"type": "fixed"}),
    ):
        node = tbl_pr.find(qn(f"w:{tag_name}"))
        if node is None:
            node = OxmlElement(f"w:{tag_name}")
            tbl_pr.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), value)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = borders.find(qn(f"w:{edge}"))
        if line is None:
            line = OxmlElement(f"w:{edge}")
            borders.append(line)
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "5")
        line.set(qn("w:space"), "0")
        line.set(qn("w:color"), BORDER)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    for tag, value in (("color", BLUE), ("u", "single"), ("sz", "22")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        r_pr.append(node)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r_pr.append(fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text, *, size=11, color=INK):
    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), size=max(8.8, size - 1), color=DARK_BLUE, mono=True)
        else:
            parsed = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if parsed:
                add_hyperlink(paragraph, parsed.group(1), parsed.group(2))
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size=size, color=color)


def add_page_number(paragraph):
    run = paragraph.add_run("")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    set_run_font(run, size=8.5, color=MUTED)


def create_decimal_numbering(doc):
    """Create a real single-level decimal list that restarts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def add_numbered_group(doc, items):
    num_id = create_decimal_numbering(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        add_inline(p, item)


def set_paragraph_border_bottom(paragraph, color=BLUE, size=12, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in (
        (1, 16, BLUE, 12, 6),
        (2, 13, BLUE, 10, 5),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        pf = style.paragraph_format
        pf.left_indent = Inches(0.50)
        pf.first_line_indent = Inches(-0.25)
        pf.space_after = Pt(8)
        pf.line_spacing = 1.167


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(header.add_run("해양 해커톤  |  병렬 딥리서치"), size=8.5, bold=True, color=MUTED)
    header.add_run("\t")
    set_run_font(header.add_run("분당독고다이  ·  2026-08-28"), size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_after = Pt(0)
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(footer.add_run("연구 완료 · 신규 제출/업로드 없음"), size=8.2, color=MUTED)
    footer.add_run("\t")
    add_page_number(footer)


def add_masthead(doc, decisions):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(kicker.add_run("DECISION MEMO  /  PARALLEL DEEP RESEARCH"), size=9, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.keep_with_next = True
    set_paragraph_border_bottom(title, color=BLUE, size=14, space=6)
    set_run_font(title.add_run("P1·P2·P3 병렬 딥리서치\n의사결정 보고서"), size=25, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(
        subtitle.add_run("제한된 공식 제출 기회를 최고점 갱신·기전 판별·손실 방어로 배분하는 실행 설계"),
        size=12,
        color=MUTED,
    )

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [1850, 7510])
    items = [
        ("작성 기준", "2026-08-28 KST"),
        ("팀", "분당독고다이"),
        ("공식 기준선", "P1 28.901363 · P2 27.264587 · P3 24.066168 · 합계 80.232118"),
        ("상태", "세 문제 독립 연구·교차검증 완료 · 신규 CSV/제출/업로드 없음"),
    ]
    for i, (label, value) in enumerate(items):
        set_cell_shading(meta.cell(i, 0), PALE_BLUE)
        p = meta.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=9.5, bold=True, color=DARK_BLUE)
        p = meta.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5, bold=(label == "상태"), color=GREEN if label == "상태" else INK)

    doc.add_paragraph()
    cards = doc.add_table(rows=1, cols=3)
    set_table_geometry(cards, [3120, 3120, 3120])
    card_data = [
        ("P1", "장구간 rescue", "최대 구조적 상한"),
        ("P2", "OAS α=.40", "즉시 exploit 1순위"),
        ("P3", "Chronos-2", "alpha축 이후 새 backbone"),
    ]
    for idx, (problem, action, note) in enumerate(card_data):
        cell = cards.cell(0, idx)
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(problem), size=9, bold=True, color=BLUE)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(action), size=14, bold=True, color=INK)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(note), size=8.5, color=MUTED)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_GOLD)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("핵심 판단  "), size=10.5, bold=True, color=GOLD)
    set_run_font(p.add_run("단일 후보의 +3점 보장은 없다. 총점 +3은 P2 저비용 수확, P1 구조 recall, P3 새 전이축을 합산해 노린다."), size=10.5)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"결정 상태: {decisions['status']}"), size=9, color=MUTED)

    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def widths_for(rows):
    cols = len(rows[0])
    presets = {
        1: [9360],
        2: [2500, 6860],
        3: [1950, 2950, 4460],
        4: [1550, 2200, 1900, 3710],
        5: [1300, 1700, 1650, 1650, 3060],
        6: [1150, 1450, 1450, 1450, 1800, 2060],
    }
    if cols in presets:
        return presets[cols]
    base = 9360 // cols
    widths = [base] * cols
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, widths_for(rows))
    repeat_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx in range(cols):
            cell = table.cell(ridx, cidx)
            if ridx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, row[cidx] if cidx < len(row) else "", size=8.4 if cols >= 5 else 8.9)
            if ridx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_inline(p, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
    if level == 1 and text.startswith(("P1", "P2", "P3")):
        p.paragraph_format.page_break_before = True


def render_markdown(doc, source_text):
    lines = source_text.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith(("부제:", "작성 기준:", "팀:", "상태:")):
            i += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = min(3, max(1, len(heading.group(1)) - 1))
            add_heading(doc, heading.group(2), level)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            items = []
            while i < len(lines):
                match = re.match(r"^(\d+)\.\s+(.+)$", lines[i].strip())
                if not match:
                    break
                items.append(match.group(2))
                i += 1
            add_numbered_group(doc, items)
            continue
        block = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("|") or re.match(r"^(-|\d+\.)\s+", nxt):
                break
            block.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(block))


def audit_docx(path):
    doc = Document(path)
    section = doc.sections[0]
    style_expectations = {
        "Normal": {"size": 11.0, "after": 6.0, "line": 1.10},
        "Heading 1": {"size": 16.0, "before": 12.0, "after": 6.0},
        "Heading 2": {"size": 13.0, "before": 10.0, "after": 5.0},
        "Heading 3": {"size": 12.0, "before": 8.0, "after": 4.0},
    }
    checks = {
        "page_width_in": round(section.page_width.inches, 3),
        "page_height_in": round(section.page_height.inches, 3),
        "margins_in": [round(x.inches, 3) for x in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)],
        "header_distance_in": round(section.header_distance.inches, 3),
        "footer_distance_in": round(section.footer_distance.inches, 3),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "table_width_sums": [],
        "styles": {},
    }
    for name, expected in style_expectations.items():
        style = doc.styles[name]
        checks["styles"][name] = {
            "font": style.font.name,
            "size": round(style.font.size.pt, 2) if style.font.size else None,
            "before": round(style.paragraph_format.space_before.pt, 2) if style.paragraph_format.space_before else 0,
            "after": round(style.paragraph_format.space_after.pt, 2) if style.paragraph_format.space_after else 0,
            "expected": expected,
        }
    for table in doc.tables:
        grid = table._tbl.tblGrid
        widths = [int(node.get(qn("w:w"))) for node in grid.findall(qn("w:gridCol"))]
        checks["table_width_sums"].append(sum(widths))

    with zipfile.ZipFile(path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        checks["has_page_field"] = " PAGE " in zf.read("word/footer1.xml").decode("utf-8")
        checks["fixed_table_layout_count"] = document_xml.count('w:tblLayout w:type="fixed"')
        checks["table_indent_120_count"] = document_xml.count('w:tblInd w:w="120" w:type="dxa"')

    pass_checks = (
        checks["page_width_in"] == 8.5
        and checks["page_height_in"] == 11.0
        and checks["margins_in"] == [1.0, 1.0, 1.0, 1.0]
        and checks["header_distance_in"] == 0.492
        and checks["footer_distance_in"] == 0.492
        and checks["has_page_field"]
        and all(value == 9360 for value in checks["table_width_sums"])
        and checks["fixed_table_layout_count"] == len(doc.tables)
        and checks["table_indent_120_count"] == len(doc.tables)
    )
    result = {"status": "PASS" if pass_checks else "FAIL", "preset": "decision_memo", "header_pattern": "memo_masthead", "checks": checks}
    AUDIT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return result


def build():
    source_text = SOURCE.read_text(encoding="utf-8")
    decisions = json.loads(DECISIONS.read_text(encoding="utf-8"))
    doc = Document()
    doc.core_properties.title = "P1·P2·P3 병렬 딥리서치 의사결정 보고서"
    doc.core_properties.subject = "해양 해커톤 제한 제출 슬롯과 구조적 돌파구"
    doc.core_properties.author = "분당독고다이"
    doc.core_properties.keywords = "P1, P2, P3, Deep Research, promotion gate, official probe"
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_masthead(doc, decisions)
    render_markdown(doc, source_text)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    result = audit_docx(OUTPUT)
    if result["status"] != "PASS":
        raise RuntimeError(json.dumps(result, ensure_ascii=False))
    print(json.dumps({"status": "PASS", "output": str(OUTPUT), "audit": str(AUDIT)}, ensure_ascii=False))


if __name__ == "__main__":
    build()
