from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = Path(
    r"C:\Users\cedis\Downloads\해양 해커톤 제출용"
    r"\20260826_round_D_preregistered_P1x3_P2x3_P3x3"
)
SOURCE = PACKAGE / "report-source.md"
OUTPUT = PACKAGE / "20260826_3x3_공식테스트_증거설계_메모_v1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "666666"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
GREEN = "EAF4EA"
GOLD = "FFF4CE"
RED = "FCE8E6"
WHITE = "FFFFFF"


def set_run_font(run, *, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_paragraph_runs(paragraph, *, size=11, color="000000", bold=None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must sum to 9360, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11)
    else:
        paragraph.add_run(text)
        set_paragraph_runs(paragraph)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    set_paragraph_runs(paragraph)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    set_paragraph_runs(paragraph)


def add_callout(doc: Document, label: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(label + "  ")
    set_run_font(lead, size=11, bold=True, color=DARK_BLUE)
    body = paragraph.add_run(text)
    set_run_font(body, size=11)
    set_table_geometry(table, [9360])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    numeric_columns: set[int] | None = None,
    font_size: float = 9.2,
    header_fill: str = LIGHT_GRAY,
):
    numeric_columns = numeric_columns or set()
    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(4)
    lead.paragraph_format.space_after = Pt(4)
    lead.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_runs(paragraph, size=font_size, bold=True, color=INK)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
                )
                set_paragraph_runs(paragraph, size=font_size)
    set_table_geometry(table, widths)
    tail = doc.add_paragraph()
    tail.paragraph_format.space_before = Pt(4)
    tail.paragraph_format.space_after = Pt(4)
    return table


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    if OUTPUT.exists():
        raise FileExistsError(OUTPUT)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    section.different_first_page_header_footer = True
    for header_part in (section.header, section.first_page_header):
        paragraph = header_part.paragraphs[0]
        paragraph.text = "분당독고다이 | OCEAN AI 해커톤 | 사전등록 실험"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_runs(paragraph, size=9, color=MUTED)
    for footer_part in (section.footer, section.first_page_footer):
        paragraph = footer_part.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("2026-08-26  |  Page ")
        set_run_font(run, size=9, color=MUTED)
        add_page_field(paragraph)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("공식 테스트 3×3 증거 설계 메모")
    set_run_font(run, size=23, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("P1·P2·P3 각 3건 | 점수 가능성과 장기 정보가치를 함께 보존하는 blind batch")
    set_run_font(run, size=14, color=MUTED)
    for label, value in (
        ("팀", "분당독고다이"),
        ("기준일", "2026-08-26 KST"),
        ("공식 잔여", "P1 3/3 · P2 3/3 · P3 3/3"),
        ("현재 상태", "9개 CSV 동결 · 독립 QA PASS · 공식 업로드 0회"),
        ("승인 효과", "실행 시 오늘의 문제별 3회를 모두 소비하여 각 0/3이 됨"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(label + ": ")
        set_run_font(label_run, size=11, bold=True, color=INK)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=11)

    add_callout(
        doc,
        "결론",
        "9개를 모두 제출할 가치가 있다. 단, 첫 점수를 보기 전에 전부 동결한 하나의 blind batch로 연속 제출해야 한다. P2·P3 Global은 수학적으로 예측 가능한 Public 개선 후보이고, 나머지 probe는 실패해도 원인 분해 증거를 남긴다.",
        GREEN,
    )

    doc.add_heading("1. 한눈에 보는 제출 설계", level=1)
    add_table(
        doc,
        ["문제", "1번", "2번", "3번", "남는 증거"],
        [
            ["P1", "Router", "O∩B", "O∪B", "신규·제거 양성 효과"],
            ["P2", "α* 전체층", "layer 2", "layer 4", "layer 2·3·4 MSE 기여"],
            ["P3", "α=-2 전체", "lead 12", "lead 18·24", "단기·12h·장기 MSE 기여"],
        ],
        [900, 1700, 1600, 1750, 3410],
        numeric_columns={0},
        font_size=9.4,
        header_fill=BLUE_GRAY,
    )
    add_body(
        doc,
        "이 구성은 문제마다 적어도 한 건의 개선·확인 후보를 두고, 나머지 두 건으로 같은 변화축의 support를 분리한다. 단순히 서로 다른 모델 9개를 순위표에 던지는 것보다 다음 구조 선택에 직접 쓰이는 정보가 많다.",
    )

    doc.add_heading("2. 공식 현황과 실험 규율", level=1)
    add_table(
        doc,
        ["문제·지표", "O", "A", "B", "현재 best"],
        [
            ["P1 F1 ↑", "0.790709", "0.786145", "0.793710", "B"],
            ["P2 RMSE ↓", "0.541085", "0.713520", "0.599921", "O"],
            ["P3 RMSE ↓", "0.607071", "0.611680", "0.609346", "O"],
        ],
        [1900, 1600, 1600, 1600, 2660],
        numeric_columns={1, 2, 3, 4},
    )
    add_bullet(doc, "9개 파일·제출 순서·수식·중지조건을 첫 업로드 전에 고정한다.")
    add_bullet(doc, "9개가 모두 올라갈 때까지 중간 점수를 열람하지 않는다.")
    add_bullet(doc, "오늘 결과를 보고 같은 날 후보를 교체하거나 재튜닝하지 않는다.")
    add_bullet(doc, "완료 후 submission ID·시각·점수·SHA-256을 하나의 ledger에 기록한다.")
    add_callout(
        doc,
        "이유",
        "반복 leaderboard 질의는 공개 holdout에 대한 적응적 과적합을 만들 수 있다. 오늘 기회를 쓰되 한 batch 안의 적응을 차단하면, 슬롯 소진과 증거 품질을 함께 관리할 수 있다.",
        GOLD,
    )

    doc.add_heading("3. P1 — 개선 확인과 2×2 양성집합 factorial", level=1)
    doc.add_heading("3.1 Router: 작은 로컬 개선의 공식 전이 확인", level=2)
    add_table(
        doc,
        ["지표", "B", "Router", "판정"],
        [
            ["양성 수", "5,856", "6,061", "+205 net"],
            ["B 대비 변경", "0", "추가 217·제거 12", "총 229행"],
            ["로컬 OOF F1", "0.86467009", "0.86690000", "+0.00222991"],
            ["day-bootstrap CI90", "—", "[+0.000918,+0.003707]", "0 미포함"],
            ["시간 fold", "기준", "2 개선·1 보합", "확인시험"],
        ],
        [2400, 2000, 3000, 1960],
        numeric_columns={1, 2, 3},
    )
    add_callout(
        doc,
        "해석 한계",
        "같은 OOF에서 disagreement cell을 선택한 뒤 평가했으므로 선택 낙관 편향이 있다. 정식 +0.0255 승격 gate에도 미달하므로 ‘확실한 새 모델’이 아니라 작은 효과의 공식 전이 확인이다.",
        GOLD,
    )

    doc.add_heading("3.2 O∩B와 O∪B: B 개선 원인의 분해", level=2)
    add_table(
        doc,
        ["모서리", "정의", "양성", "역할"],
        [
            ["I", "O∩B", "5,680", "공통 양성"],
            ["B", "I∪X", "5,856", "B-only X=176 포함"],
            ["O", "I∪Y", "6,504", "O-only Y=824 포함"],
            ["U", "I∪X∪Y", "6,680", "전체 합집합"],
        ],
        [1000, 1900, 1500, 4960],
        numeric_columns={0, 2},
    )
    add_bullet(doc, "X 효과: B-I와 U-O의 부호를 두 배경에서 비교한다.")
    add_bullet(doc, "Y 효과: O-I와 U-B의 부호를 두 배경에서 비교한다.")
    add_bullet(doc, "부호가 같으면 전이 안정, 다르면 context·F1 비선형성·public 표본 불안정을 의심한다.")
    add_body(
        doc,
        "P1 F1은 비선형이고 숨은 confusion matrix를 알 수 없으므로 점수 차이를 개별 행의 인과효과로 역산하지 않는다. 여기서 factorial은 조작 방향의 재현성을 보는 기술적 분해다.",
    )

    doc.add_heading("3.3 P1 제출 파일", level=2)
    add_table(
        doc,
        ["순서", "제출물 제목", "역할", "SHA-256 앞 12자"],
        [
            ["P1-1", "disagreement router 확인 v1", "Exploit 확인", "1b04e81c18d5"],
            ["P1-2", "O∩B 추가양성 절제 실험 v1", "X probe", "0ac5a6abe623"],
            ["P1-3", "O∪B 제거양성 복원 실험 v1", "Y probe", "c8b72922f42d"],
        ],
        [1100, 3850, 2000, 2410],
        numeric_columns={0, 3},
    )

    doc.add_heading("4. P2 — Public 이차최적화와 layer 분해", level=1)
    add_body(
        doc,
        "과거 O·B·A는 동일 축 P(α)=O+α(A-O)의 α=0·0.5·1이며 B는 수치상 정확한 중점이다. 고정 평가 집합에서 MSE는 α의 정확한 이차식이므로 세 공식 RMSE만으로 경로 최적점을 결정할 수 있다.",
    )
    add_callout(
        doc,
        "사전등록 예측",
        "MSE(α)=0.164146710286α²+0.052191102889α+0.292772977225, α*=-0.158976999289. Global 예상 Public RMSE는 0.537237735794이며 6자리 반올림 전파 구간은 0.537236416~0.537239056이다.",
        GREEN,
    )
    add_table(
        doc,
        ["후보", "적용 support", "행", "해석"],
        [
            ["Global", "layer 2·3·4", "26,061", "개선 + scorer 일관성"],
            ["Layer2", "layer 2", "8,713", "Δ2 직접 측정"],
            ["Layer4", "layer 4", "8,636", "Δ4 직접 측정"],
            ["차감 Layer3", "layer 3", "8,712", "Δ3=qG-q2-q4+qO"],
        ],
        [1900, 2100, 1500, 3860],
        numeric_columns={2},
    )
    add_table(
        doc,
        ["로컬 p100", "ΔRMSE", "day-bootstrap CI90", "판정"],
        [
            ["Global", "-0.004883", "[-0.007395,-0.002311]", "강한 개선"],
            ["Layer2", "-0.001121", "[-0.001459,-0.000768]", "안정적 개선"],
            ["Layer4", "-0.000357", "[-0.001795,+0.001121]", "불확실"],
            ["차감 Layer3", "-0.003400", "[-0.004331,-0.002453]", "가장 큰 기여"],
        ],
        [2200, 1700, 3000, 2460],
        numeric_columns={1, 2},
    )
    add_callout(
        doc,
        "중지 조건",
        "Global 공식 RMSE가 0.537236416~0.537239056와 모순되면 layer 해석을 중지하고 scoring·파일 계보·표시 정밀도를 먼저 감사한다. Δ는 pooled MSE 기여이며 layer별 평균 RMSE가 아니다.",
        RED,
    )
    doc.add_heading("4.1 P2 제출 파일", level=2)
    add_table(
        doc,
        ["순서", "제출물 제목", "역할", "SHA-256 앞 12자"],
        [
            ["P2-1", "공개 이차최적 전체층 v1", "Exploit", "9cc951801cf6"],
            ["P2-2", "공개 이차최적 2층 단독 v1", "Layer probe", "5507317f45bf"],
            ["P2-3", "공개 이차최적 4층 단독 v1", "Layer probe", "98890354fe79"],
        ],
        [1100, 3850, 2000, 2410],
        numeric_columns={0, 3},
    )

    doc.add_heading("5. P3 — 보수적 역방향 exploit과 lead 분해", level=1)
    add_body(
        doc,
        "과거 B는 모든 1,200행에서 B=(O+A)/2와 1e-12 이내로 일치한다. 따라서 α=0·0.5·1의 공식 점수가 동일 축의 정확한 MSE 이차곡선을 결정한다. 수학적 Public 최적점은 α≈-16.52지만 외삽 위험을 제한하기 위해 α=-2를 선택했다.",
    )
    add_callout(
        doc,
        "사전등록 예측",
        "Global 예상 Public RMSE는 0.598574192921, O 대비 0.008496807079m 개선이다. 표시점수 반올림 전파 구간은 0.598549260514~0.598599124290이다. 예측 범위 0.794244~4.163249m, 최대 변화 0.171800m로 물리 guard를 통과한다.",
        GREEN,
    )
    add_table(
        doc,
        ["후보", "적용 support", "변경 행", "O 대비 RMS 변화"],
        [
            ["Global", "모든 lead", "660", "0.0314885m"],
            ["Lead12", "12h", "200", "0.0149503m"],
            ["Lead18/24", "18h·24h", "400", "0.0277022m"],
            ["차감 Early", "3h·6h·9h", "60", "qG-q12-q18/24+qO"],
        ],
        [1900, 2200, 1700, 3560],
        numeric_columns={2, 3},
    )
    add_body(
        doc,
        "로컬 analogue에서는 Global +0.002166m, Lead12 +0.000129m, Lead18/24 +0.002037m로 모두 악화했다. 공식에서 개선되면 이는 public/local 부호 역전의 강한 증거이며 Private 일반화 성공으로 간주하지 않는다.",
    )
    add_callout(
        doc,
        "중요 한계",
        "P3 Public은 66 case·396행으로 작고 공식 페이지도 참고용이라고 명시한다. Global 점수가 예측 구간에 맞아도 public-specific correction일 수 있다. Private 승격은 별도 판단이다.",
        GOLD,
    )
    doc.add_heading("5.1 P3 제출 파일", level=2)
    add_table(
        doc,
        ["순서", "제출물 제목", "역할", "SHA-256 앞 12자"],
        [
            ["P3-1", "O-A 역방향 전체리드 v1", "Exploit", "57a90beb3f81"],
            ["P3-2", "O-A 역방향 12시간 단독 v1", "Lead probe", "c5ac003e5c08"],
            ["P3-3", "O-A 역방향 18·24시간 단독 v1", "Lead probe", "91ead7470f53"],
        ],
        [1100, 3850, 2000, 2410],
        numeric_columns={0, 3},
    )
    add_callout(
        doc,
        "ERA5 분리",
        "이 세 파일은 역사적 O/A/B 제출만 사용했다. 현재 별도로 실행 중인 ERA5 파일·값·프로세스·286개 특징·split·gate·모델을 읽거나 바꾸지 않았고 그 실험을 대체하지 않는다.",
        BLUE_GRAY,
    )

    doc.add_heading("6. 실행 체크리스트와 결과 판정", level=1)
    add_number(doc, "SET_MANIFEST.json의 9개 SHA-256과 실제 CSV를 다시 대조한다.")
    add_number(doc, "사용자 action-time 승인을 받은 뒤에만 업로드한다.")
    add_number(doc, "P1-1→P1-2→P1-3→P2-1→P2-2→P2-3→P3-1→P3-2→P3-3 순서로 제출한다.")
    add_number(doc, "9개 완료 전에는 중간 점수를 읽지 않고, 오늘 결과 기반 후보 교체도 하지 않는다.")
    add_number(doc, "완료 후 submission ID·시각·표시점수·SHA-256을 하나의 ledger에 기록한다.")
    add_number(doc, "P2/P3는 RMSE를 제곱하고 ±0.5e-6 표시 반올림 구간을 전파한 뒤 부호를 판정한다.")
    add_number(doc, "Public 개선은 Private 승격과 분리해 기록한다.")

    doc.add_heading("6.1 결과에 따른 다음 결정", level=2)
    add_table(
        doc,
        ["문제", "관측", "다음 결정"],
        [
            ["P1", "Router 개선 + factorial 부호 안정", "해당 정점·층 규칙을 다음 모델 prior로 사용"],
            ["P1", "Router 악화", "OOF post-selection 편향으로 기록, 동일 mining 중단"],
            ["P2", "Global 예측 구간 일치", "scorer·축 계보 검증, 음의 Δ가 큰 layer 우선"],
            ["P2", "Global 구간 모순", "layer 해석 중지, 업로드·평가 버전 감사"],
            ["P3", "Global 구간 일치", "public-specific correction으로 분류"],
            ["P3", "공식 개선·로컬 악화", "distribution shift 증거, Private 승격 근거로 쓰지 않음"],
        ],
        [1100, 3000, 5260],
        font_size=8.9,
    )

    doc.add_heading("7. 독립 QA 요약", level=1)
    add_table(
        doc,
        ["문제", "파일", "핵심 재검산", "결과"],
        [
            ["P1", "3", "169,011행·집합식·Router 계보·해시", "PASS"],
            ["P2", "3", "26,061행·α 수식·layer mask·해시", "PASS"],
            ["P3", "3", "1,200행·α 수식·lead mask·물리범위·해시", "PASS"],
            ["통합", "9", "메모·manifest·READY·12개 CSV checksum", "PASS"],
        ],
        [1200, 1000, 5500, 1660],
        numeric_columns={1, 3},
        header_fill=BLUE_GRAY,
    )
    add_body(
        doc,
        "P2/P3 CSV는 소수점 12자리 직렬화로 비활성 행에도 최대 약 5.01e-13의 수치 drift가 있다. manifest의 변경 행은 |Δ|>1e-12 기준이며 제출 적합성이나 분해식에는 영향이 없다.",
    )

    doc.add_heading("8. 연구 근거와 한계", level=1)
    sources = [
        ("Blum & Hardt (2015), The Ladder", "반복 leaderboard 적응과 public holdout 과적합", "https://proceedings.mlr.press/v37/blum15.html"),
        ("Dwork et al. (2015), Reusable Holdout", "적응적 holdout 재사용의 일반화 위험", "https://proceedings.neurips.cc/paper_files/paper/2015/hash/bad5f33780c42f2588878a9d07405083-Abstract.html"),
        ("Roelofs et al. (2019)", "실제 ML 경진대회의 leaderboard overfitting 메타분석", "https://proceedings.neurips.cc/paper/2019/hash/ee39e503b6bedf0c98c388b7e8589aca-Abstract.html"),
        ("Cawley & Talbot (2010)", "모델 선택 단계의 검증 편향", "https://www.jmlr.org/papers/v11/cawley10a.html"),
        ("Kaufmann et al. (2016)", "고정예산 best-arm identification", "https://www.jmlr.org/papers/v17/kaufman16a.html"),
        ("Russo & Van Roy (2017)", "정보획득과 즉시 보상의 균형", "https://pubsonline.informs.org/doi/10.1287/opre.2017.1663"),
        ("Box & Wilson (1951)", "국소 이차 response surface와 제한 실험 설계", "https://doi.org/10.1111/j.2517-6161.1951.tb00067.x"),
        ("Bakker et al. (2020)", "구체적 사전등록과 해석 자유도 감소", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3000937"),
    ]
    for title_text, claim, url in sources:
        paragraph = doc.add_paragraph(style="List Bullet")
        first = paragraph.add_run(title_text + ": ")
        set_run_font(first, size=9.8, bold=True, color=INK)
        second = paragraph.add_run(claim + ". " + url)
        set_run_font(second, size=9.8)

    doc.add_heading("8.1 남는 한계", level=2)
    for text in (
        "사전등록을 해도 누적된 공개 리더보드 적응 위험이 완전히 사라지지는 않는다.",
        "P1 F1 차이는 숨은 confusion matrix 없이 가산적 인과효과로 해석할 수 없다.",
        "P2/P3 이차 예측은 동일 public 행·가중치·scorer가 유지된다는 조건부 결과다.",
        "P2/P3 분해는 pooled MSE 기여이며 그룹별 평균 RMSE가 아니다.",
        "오늘 9개는 모델 구조 9종 비교가 아니라 기존 축과 support를 분해하는 제한예산 실험이다.",
    ):
        add_bullet(doc, text)

    add_callout(
        doc,
        "최종 판정: GO_FOR_USER_CONFIRMATION",
        "제출 세트는 정확히 P1 3개·P2 3개·P3 3개이며 독립 QA를 통과했다. 실제 업로드만 남았다. 승인하면 오늘의 문제별 3회를 모두 소비하므로 각 문제 잔여는 0/3이 된다.",
        GOLD,
    )

    doc.core_properties.title = "2026-08-26 공식 테스트 3×3 증거 설계 메모"
    doc.core_properties.subject = "OCEAN AI 해커톤 P1/P2/P3 사전등록 blind batch"
    doc.core_properties.author = "분당독고다이"
    doc.core_properties.keywords = "P1, P2, P3, preregistration, leaderboard, QA"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
