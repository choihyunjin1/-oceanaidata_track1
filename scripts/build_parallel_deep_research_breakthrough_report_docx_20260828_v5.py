"""Build the Korean P1/P2/P3 breakthrough research decision brief v5."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import build_parallel_deep_research_execution_report_docx_20260828_v4 as base


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = (
    ROOT
    / "reports"
    / "parallel_deep_research_breakthrough_20260828_v5"
    / "P1_P2_P3_새_돌파구_딥리서치_결론보고_v5_20260828.docx"
)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    base.set_run_font(kicker.add_run("TECHNICAL DECISION BRIEF"), size=9.5, bold=True, color=base.BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    base.set_run_font(title.add_run("P1·P2·P3 새 돌파구 딥리서치 v5"), size=23, bold=True, color="000000")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    base.set_run_font(
        subtitle.add_run("과거 실패 계열 제외, 1차 출처 검증, 구조 후보 3개, 다음 실행 사전등록"),
        size=12.5,
        color=base.MUTED,
    )

    rows = (
        ("팀", "분당독고다이"),
        ("작성일", "2026-08-28 KST"),
        ("문서 성격", "내부 연구 의사결정 및 실행 전 사전등록"),
        ("안전 경계", "공식 test/sample/submission 접근·신규 CSV·업로드 모두 0회"),
    )
    for label, value in rows:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        base.set_run_font(paragraph.add_run(f"{label}: "), size=10.5, bold=True, color="000000")
        base.set_run_font(paragraph.add_run(value), size=10.5, color="000000")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(9)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), base.BLUE)
    borders.append(bottom)
    rule._p.get_or_add_pPr().append(borders)


def configure_page(doc: Document) -> None:
    base.configure_page(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    for run in list(header.runs):
        run._element.getparent().remove(run._element)
    base.set_run_font(header.add_run("분당독고다이 | 새 돌파구 딥리서치 v5"), size=8.5, color=base.MUTED)


def build(output: Path) -> None:
    doc = Document()
    base.configure_styles(doc)
    configure_page(doc)
    add_title_block(doc)

    base.add_callout(
        doc,
        "결론",
        "새 구조 1순위는 P2 BayOTIDE형 동적 저차원 상태공간, 2순위는 P1 TS2Vec형 normal-prototype proposal generator, 3순위는 P3 TimeXer형 past-exogenous direct 6-lead다. 별도로 P2 OAS α=.40은 이미 QA된 공식 probe이므로 새 구조 연구와 분리해 판단한다.",
        fill=base.PALE_GREEN,
    )
    base.add_table(
        doc,
        ["우선", "문제", "구조", "직접 겨냥한 병목", "상태"],
        [
            ["1", "P2", "BayOTIDE형 dynamic SSM", "정적 profile·점별 보정의 시간동역학 부족", "사전등록"],
            ["2", "P1", "TS2Vec normal prototype", "양성 event support 부족·proposal bank 고정", "사전등록"],
            ["3", "P3", "TimeXer direct 6-lead", "expert OOF support 1/3·router coverage 부족", "사전등록"],
            ["별도", "P2", "Seasonal OAS α=.40", "공식 α=.10→.20 상승 축의 정보가치", "Probe ready"],
        ],
        [720, 720, 2440, 3780, 1700],
    )
    base.add_body(
        doc,
        "문헌 benchmark의 개선률을 대회 점수로 옮기지 않았다. 공식 종합점수와 로컬 F1·°C·m는 별도 척도이며, 새 구조는 모두 실행 미승인 상태다.",
        bold_lead="해석 원칙. ",
    )

    doc.add_heading("1. 연구 질문과 제외선", level=1)
    base.add_body(
        doc,
        "이번 회차의 질문은 ‘기존 family의 threshold를 조금 더 맞출 수 있는가’가 아니라 ‘확인된 실패 원인을 다른 정보경로로 우회할 수 있는가’였다. 각 문제를 독립 조사하고 최고 영향 주장을 원 논문과 저자 공식 코드에서 다시 확인했다."
    )
    base.add_table(
        doc,
        ["문제", "다시 하지 않는 것", "새 정보경로"],
        [
            ["P1", "83-bank ranker·masked reconstruction·interval threshold 구조 구제", "라벨 없는 hierarchical contrastive representation"],
            ["P2", "1-mode heave·rank-4 CMFPCA·safe gate 재탐색", "공개 T/S가 계속 갱신하는 dynamic latent state"],
            ["P3", "ERA5 router threshold·Chronos step·TSMixer weight 재탐색", "내생 파고와 과거 외생 기상의 비대칭 direct model"],
        ],
        [950, 4010, 4400],
    )

    doc.add_heading("2. P2 1순위 — BayOTIDE형 동적 공동 복원", level=1)
    base.add_body(
        doc,
        "BayOTIDE는 다변량 시계열을 서로 다른 pattern의 functional low-rank factor로 분해하고 GP prior를 equivalent SDE/state-space로 바꿔 전역 trend와 주기구조를 갱신한다. P2에서는 actual-depth로 정렬한 T/S 채널 전체가 대상이며, target layer 2·3·4의 T/S는 결측 블록에서 fit과 update 모두에서 배제한다."
    )
    base.add_table(
        doc,
        ["항목", "사전등록 값", "이유"],
        [
            ["Latent factors", "Matérn-3/2 trend 3 + periodic 12.42h·24h 각 1", "물리 주기를 포함하되 grid search 0회"],
            ["Inference", "Forward filter + RTS smoother", "가림 전후와 공개층 문맥을 시간상태로 결합"],
            ["Primary mask", "L2-L4 T/S 동시 block mask", "실제 결측 메커니즘 모사"],
            ["Comparator", "p2_extrapolated_soft_gate_v2 common OOF", "가장 강한 공통 비교면 유지"],
            ["승격", "Δ≤-0.003°C, CI90 upper<0, 2/3 folds, worst layer≤+0.005°C", "작은 fitted 이득과 층 붕괴 동시 차단"],
        ],
        [1800, 3970, 3590],
    )
    base.add_callout(
        doc,
        "가장 큰 반증",
        "논문의 real-data missingness는 P2의 61일 target-channel blackout과 동일하지 않다. Random point mask 성능은 승격 근거로 쓰지 않고, 기존 세 historical block과 7일 purge에서만 판정한다.",
    )
    base.add_body(
        doc,
        "공식 OAS 10%→20%의 상승은 수직 T/S 공분산 family가 숨은 평가면에서 유효할 가능성을 보여줬다. BayOTIDE는 그 신호를 단순 blend 확대가 아니라 시간에 따라 변화하는 latent profile로 확장한다.",
        bold_lead="우선순위 근거. ",
    )

    doc.add_heading("3. P1 2순위 — TS2Vec 조건부 normal prototype", level=1)
    base.add_body(
        doc,
        "P1의 병목은 event verifier가 약해서가 아니라 독립 qualification에 양성 proposal이 1개뿐이어서 학습 자체가 성립하지 않은 점이다. TS2Vec는 augmented context view 사이의 hierarchical contrastive agreement로 timestamp representation을 학습하므로 anomaly label 없이 proposal support를 새로 만들 수 있다."
    )
    base.add_table(
        doc,
        ["항목", "사전등록 값", "보호장치"],
        [
            ["Encoder fit", "2024-09-15 이전, label-free", "Loss·sampling에 label 미사용"],
            ["Window", "512 rows, multiscale timestamp embedding", "연속 segment gap 횡단 금지"],
            ["Normal score", "Station×layer seasonal prototype + global shrinkage + kNN", "Test prevalence·SPOT·point adjustment 금지"],
            ["Candidate", "Frozen e150 anchor OR additions", "Anchor 1 삭제 금지"],
            ["승격", "두 구간 ΔF1>0, new event≥2, ≥2 cells, bootstrap P≥.8", "FP/day cap 동시 통과"],
        ],
        [1800, 3910, 3650],
    )
    base.add_callout(
        doc,
        "가장 큰 반증",
        "TS2Vec의 published anomaly protocol은 point anomaly와 local adjustment 중심이다. 우리는 encoder만 차용하고 P1 event decoder와 honest historical qualification으로 다시 검증한다.",
    )

    doc.add_heading("4. P3 3순위 — TimeXer형 past-exogenous direct 6-lead", level=1)
    base.add_body(
        doc,
        "TimeXer는 endogenous temporal patch self-attention과 exogenous variate cross-attention을 분리한다. P3에서는 과거 48시간 hs만 endogenous로, tp·hmax·wind·pressure·air state·mask만 exogenous로 사용하고 미래 기상·절대시각은 금지한다. 출력은 6개 lead residual의 직접 joint head다."
    )
    base.add_table(
        doc,
        ["항목", "사전등록 값", "보호장치"],
        [
            ["출력", "3/6/9/12/18/24h direct residual", "Lead별 사후 model 선택 금지"],
            ["Seeds", "142857 / 271828 / 314159", "Inner-best checkpoint ensemble"],
            ["검증", "3 outer, prediction seal 후 truth attach", "Expert router·future covariate 미사용"],
            ["승격", "Δ≤-0.005m, CI90 upper<0, 2/3 folds", "Worst station≤+0.01m"],
            ["장리드", "12/18/24h 모두 비악화", "P3 총점의 long-lead 붕괴 차단"],
        ],
        [1800, 3910, 3650],
    )
    base.add_callout(
        doc,
        "가장 큰 반증",
        "3개 station과 제한된 anchor에서 Transformer는 쉽게 과적합한다. Inner에서 persistence를 안정적으로 이기지 못하면 outer를 실행하지 않고 즉시 중단한다.",
    )

    doc.add_heading("5. 즉시 가치가 있는 별도 축 — P2 OAS α=.40", level=1)
    base.add_body(
        doc,
        "OAS α=.40은 이번 문헌조사에서 새로 만든 모델이 아니다. 이미 공식 α=.10과 α=.20이 연속 개선됐고, α=.40 파일은 lineage reproduction, 26,061행 schema, PAVA idempotence, SHA-256 독립 QA를 통과했다."
    )
    base.add_table(
        doc,
        ["항목", "확인값", "의미"],
        [
            ["공식 α=.10", "RMSE 0.507628 / 26.963865점", "첫 구조 probe 개선"],
            ["공식 α=.20", "RMSE 0.483661 / 27.264587점", "같은 방향 추가 개선"],
            ["α=.40 조건부 예상", "RMSE 중심 0.448627", "보장은 아닌 기하 기반 추정"],
            ["α=.20 대비 기대", "약 +0.23~+0.44점", "공식 한 슬롯의 정보가치"],
            ["현재 권한", "업로드 미승인", "정확한 파일·SHA 재승인 필요"],
        ],
        [2100, 2960, 4300],
    )
    base.add_callout(
        doc,
        "의사결정 분리",
        "OAS40은 단기 공식 probe이고 BayOTIDE·TS2Vec·TimeXer는 새 구조 연구다. OAS40을 제출하더라도 새 구조의 우열이 증명되는 것은 아니며, 로컬 OOF가 반대로 악화한 사실도 함께 기록한다.",
        fill=base.PALE_BLUE,
    )

    doc.add_heading("6. 실행 순서와 중단선", level=1)
    base.add_table(
        doc,
        ["순서", "작업", "예상 자원", "첫 중단선"],
        [
            ["A", "P2 OAS40 공식 probe 여부 결정", "업로드 1회", "정확한 파일 승인 없으면 0회"],
            ["1", "P2 BayOTIDE bounded run", "낮음~중간, 약 1~3h", "Mask/leakage/OOF gate 실패"],
            ["2", "P1 TS2Vec bounded screen", "RTX5090 약 30~90m", "Embedding coverage·FP/day·event support 실패"],
            ["3", "P3 TimeXer nested run", "RTX5090 약 2~4h", "Inner persistence 안정성 실패"],
        ],
        [780, 3090, 2040, 3450],
    )
    base.add_body(
        doc,
        "결과를 본 뒤 factor 수, context/window, epoch, threshold, patch width를 낮추거나 늘려 같은 family를 구제하지 않는다. Best checkpoint는 마지막 epoch가 아니라 사전등록 inner-best로 복원한다.",
        bold_lead="재실행 금지선. ",
    )

    doc.add_heading("7. 근거와 불확실성", level=1)
    base.add_table(
        doc,
        ["ID", "1차 출처", "검증한 주장", "남은 gap"],
        [
            ["P1-1", "Yue et al., AAAI 2022 / official code", "Hierarchical contrastive timestamp representation", "P1 long-event F1"],
            ["P2-1", "Fang et al., ICML 2024 / official code", "Functional low-rank GP→SDE state-space", "61-day block mask"],
            ["P2-2", "Nie et al., KDD 2024 / official code", "Low-rank attention block imputation", "Actual-depth 변형"],
            ["P3-1", "Wang et al., NeurIPS 2024 / official code", "Endogenous/exogenous asymmetric attention", "Past-only small sample"],
            ["P3-2", "SWAN technical documentation", "Wave action/energy source-sink 원리", "Bulk point data만으로 불완전"],
        ],
        [850, 3120, 2970, 2420],
    )
    base.add_body(
        doc,
        "상세 URL, contradiction, confidence, remaining gap은 같은 폴더의 claim-ledger.json에 있다. 공개 코드가 있다는 사실은 대회 데이터에서 성능이 난다는 증거가 아니다."
    )

    base.add_callout(
        doc,
        "8. 운영 규칙 업데이트 — 마감 충돌 발견",
        "2026-08-28 공개 공식 홈페이지는 대학부 예선 종료·결과물 제출 마감을 2026-09-30으로 표시한다. 저장소의 9월 7일 메모와 충돌하므로 최종모델 잠금·제출 같은 행동 전 로그인 공지와 문제 상세를 다시 확인해야 한다.",
        fill=base.PALE_RED,
    )

    base.add_callout(
        doc,
        "9. 최종 판정",
        "전체 모델 공간을 완전히 탐색한 것은 아니다. 다만 과거 실패 계열과 겹치지 않고, 1차 근거·공식 코드·현재 컴퓨터에서의 bounded 검증 가능성을 갖춘 세 구조로 수렴했다. 실행 우선순위는 BayOTIDE → TS2Vec → TimeXer다. P2 OAS40 공식 probe는 별도 승인이 필요하다.",
        fill=base.PALE_GREEN,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
