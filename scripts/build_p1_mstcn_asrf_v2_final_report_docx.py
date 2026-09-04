"""Create the visually verified P1 MS-TCN++/ASRF v2 research brief.

The Markdown research report is the canonical narrative source.  This builder
adds a deterministic `standard_business_brief` layout using the
`memo_masthead` first-page pattern and embeds the three verified evidence
figures from the post-run bundle.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "p1_incumbent_preserving_mstcn_asrf_v2"
SOURCE = REPORT_DIR / "report-source.md"
FIG_DIR = REPORT_DIR / "postrun_bundle" / "figures"
OUTPUT = REPORT_DIR / "20260827_P1_MSTCN_ASRF_v2_NO_GO_기술보고서.docx"

# standard_business_brief tokens
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "596574"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_RED = "FDECEC"
PALE_GREEN = "E8F5EC"
GREEN = "276749"
RED = "9B1C1C"
GOLD = "7A5A00"
BORDER = "D6DCE5"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, *, size=11, bold=None, italic=None, color=INK, mono=False):
    latin = "Consolas" if mono else "Calibri"
    east = "Malgun Gothic"
    run.font.name = latin
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in CELL_MARGINS.items():
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    props = table._tbl.tblPr
    for tag, attrs in (
        ("tblW", {"w": str(CONTENT_DXA), "type": "dxa"}),
        ("tblInd", {"w": str(TABLE_INDENT_DXA), "type": "dxa"}),
        ("tblLayout", {"type": "fixed"}),
    ):
        node = props.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            props.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    borders = props.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        props.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = borders.find(qn(f"w:{edge}"))
        if line is None:
            line = OxmlElement(f"w:{edge}")
            borders.append(line)
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "5")
        line.set(qn("w:space"), "0")
        line.set(qn("w:color"), BORDER)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_hyperlink(paragraph, text, url):
    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    for tag, value in (("color", BLUE), ("u", "single"), ("sz", "21")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        rpr.append(node)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rpr.append(fonts)
    run.append(rpr)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(txt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text, *, size=11, color=INK):
    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), size=max(8.5, size - 1), color=DARK_BLUE, mono=True)
        else:
            parsed = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            label, target = parsed.group(1), parsed.group(2)
            if target.startswith(("http://", "https://")):
                add_hyperlink(paragraph, label, target)
            else:
                set_run_font(paragraph.add_run(label), size=size, bold=True, color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size=size, color=color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.keep_with_next = True
        pf.keep_together = True
    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True
    list_bullet = doc.styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    list_bullet.font.size = Pt(11)
    list_bullet.paragraph_format.left_indent = Inches(0.5)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.25)
    list_bullet.paragraph_format.space_after = Pt(8)
    list_bullet.paragraph_format.line_spacing = 1.167


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("P1 구조 연구  |  MS-TCN++/ASRF v2"), size=8.5, bold=True, color=MUTED)
    header.add_run("\t")
    set_run_font(header.add_run("분당독고다이 · 2026-08-27"), size=8.5, color=MUTED)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("내부 연구 증거 · 공식 제출 후보 아님"), size=8.2, color=MUTED)
    footer.add_run("\t")
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    add_page_field(footer)


def add_bottom_rule(paragraph, color=BLUE, size="12"):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_masthead(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(kicker.add_run("P1 HIGH-CAPACITY STRUCTURAL RESEARCH"), size=9, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("Incumbent-preserving\nMS-TCN++/ASRF v2"), size=25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("고용량 장기 사건·경계 복구 모델의 사전등록 실행, 독립 QA 및 승격 판정"), size=12, color=MUTED)
    metadata = [
        ("팀", "분당독고다이"),
        ("실행", "2026-08-27 05:47:41–10:03:14 KST · 4시간 15분 33초"),
        ("판정", "NO_GO_CONFIRMATORY · 공식 제출/업로드 0회"),
        ("계약", "width 512 · epoch 125 · threshold 0.9 · 3-seed raw mean"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(f"{label}: "), size=10.5, bold=True, color=DARK_BLUE)
        set_run_font(p.add_run(value), size=10.5, bold=(label == "판정"), color=RED if label == "판정" else INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    add_bottom_rule(rule)

    decision = doc.add_table(rows=1, cols=1)
    set_table_geometry(decision, [CONTENT_DXA])
    cell = decision.cell(0, 0)
    shade(cell, PALE_RED)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("결정  |  공식 제출 후보 기각"), size=12, bold=True, color=RED)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run("Q2의 +0.098157은 확인창에서 재현되지 않았다. Q4 false positive가 급증해 pooled ΔF1은 -0.005140이며, 새 전략은 다중 창 최악 성능과 추가행 precision을 직접 최적화해야 한다."),
        size=10.5,
        color=INK,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    cards = doc.add_table(rows=1, cols=4)
    set_table_geometry(cards, [2340, 2340, 2340, 2340])
    values = [
        ("Pooled ΔF1", "-0.005140", RED),
        ("Q4 ΔF1", "-0.031484", RED),
        ("추가 precision", "38.11%", GOLD),
        ("P(ΔF1>0)", "35.15%", GOLD),
    ]
    for idx, (label, value, color) in enumerate(values):
        cell = cards.cell(0, idx)
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(label), size=8.5, bold=True, color=MUTED)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=15, bold=True, color=color)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(note.add_run("독립 QA: 40/40 산출물 · 137 assertions PASS · 지표/게이트/bootstrap 재현"), size=9.5, bold=True, color=GREEN)
    doc.add_page_break()


def ensure_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1
    abstract_ids = {}
    abstract_nodes = []
    # Word/LibreOffice's portable bullet encoding is U+F0B7 rendered through
    # Symbol.  A plain U+2022 is treated as a decimal fallback by LibreOffice.
    for kind, fmt, text in (("bullet", "bullet", "\uf0b7"), ("decimal", "decimal", "%1.")):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        if kind == "bullet":
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), "Symbol")
            rfonts.set(qn("w:hAnsi"), "Symbol")
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract_ids[kind] = next_abs
        abstract_nodes.append(abstract)
        next_abs += 1
    # NumberingML requires all abstractNum elements to precede concrete num
    # elements.  Inserting at the first concrete num keeps the package valid.
    first_num = numbering.find(qn("w:num"))
    insert_at = list(numbering).index(first_num) if first_num is not None else len(numbering)
    for node in abstract_nodes:
        numbering.insert(insert_at, node)
        insert_at += 1
    return {"numbering": numbering, "abstract_ids": abstract_ids, "next_num": next_num}


def new_num_id(manager, kind):
    num_id = manager["next_num"]
    manager["next_num"] += 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(manager["abstract_ids"][kind]))
    num.append(abs_id)
    if kind == "decimal":
        # A distinct numId is not enough to restart numbering in LibreOffice;
        # carry an explicit level start override for each numbered sequence.
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), "1")
        override.append(start)
        num.append(override)
    manager["numbering"].append(num)
    return num_id


def add_list_item(doc, text, num_id=None, *, kind="decimal"):
    p = doc.add_paragraph(style="List Bullet" if kind == "bullet" else None)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    if kind == "decimal":
        ppr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, num))
        ppr.append(num_pr)
    add_inline(p, text)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        rows.append([c.strip() for c in lines[idx].strip().strip("|").split("|")])
        idx += 1
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, idx


def table_widths(cols):
    return {
        1: [9360],
        2: [2500, 6860],
        3: [2500, 2500, 4360],
        4: [1900, 2400, 2200, 2860],
        5: [1300, 1800, 1800, 1800, 2660],
        6: [1300, 1900, 1900, 1500, 1300, 1760],
        7: [1250, 1580, 1580, 1200, 1200, 1200, 1350],
    }.get(cols, [CONTENT_DXA // cols] * (cols - 1) + [CONTENT_DXA - (CONTENT_DXA // cols) * (cols - 1)])


def add_data_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, table_widths(cols))
    repeat_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx in range(cols):
            cell = table.cell(ridx, cidx)
            if ridx == 0:
                shade(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, row[cidx] if cidx < len(row) else "", size=8.1 if cols >= 6 else 8.7)
            if ridx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif cidx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figure(doc, filename, caption, alt_text):
    path = FIG_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.keep_with_next = True
    set_run_font(cap.add_run(caption), size=9, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.2))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)


def render_markdown(doc, text, numbering_manager):
    lines = text.splitlines()
    idx = 0
    decimal_id = None
    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            idx += 1
            continue
        if stripped.startswith("# ") or stripped.startswith("기술 보고서") or stripped.startswith("팀:") or stripped.startswith("상태:"):
            idx += 1
            continue
        heading = re.match(r"^(#{2,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            title = heading.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, title, size=16 if level == 1 else 13, color=BLUE if level <= 2 else DARK_BLUE)
            idx += 1
            if title == "실행 결과":
                add_figure(
                    doc,
                    "figure_01_training_loss_convergence.png",
                    "그림 1. 12개 학습 이력의 손실 곡선과 후반 불안정성",
                    "Q2, Q3, Q4 학습 손실 곡선. 모든 값은 유한하지만 일부 seed의 후반 손실이 재상승한다.",
                )
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_data_table(doc, rows)
            if rows and rows[0] and rows[0][0] == "기준" and any("long-event" in c for row in rows for c in row):
                add_figure(
                    doc,
                    "figure_02_q2_qualification_envelope.png",
                    "그림 2. Q2 사양 선택 envelope와 epoch 125 고립 peak",
                    "Width 256과 512의 Q2 후보 envelope. 선택 epoch 125는 인접 epoch보다 높은 고립 peak다.",
                )
            if rows and rows[0] and rows[0][0] == "Station":
                add_figure(
                    doc,
                    "figure_03_confirmatory_effects_and_gates.png",
                    "그림 3. 확인창·station별 효과와 사전 고정 gate",
                    "Q3는 개선됐지만 Q4와 두 station이 악화되어 pooled 및 high-impact gate가 실패했다.",
                )
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            add_list_item(doc, bullet.group(1), kind="bullet")
            idx += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            if decimal_id is None or numbered.group(1) == "1":
                decimal_id = new_num_id(numbering_manager, "decimal")
            add_list_item(doc, numbered.group(2), decimal_id)
            idx += 1
            continue
        block = [stripped]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("|") or re.match(r"^(-|\d+\.)\s+", nxt):
                break
            block.append(nxt)
            idx += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(block))


def add_sources(doc):
    p = doc.add_paragraph(style="Heading 1")
    add_inline(p, "1차 문헌·구현 근거", size=16, color=BLUE)
    intro = doc.add_paragraph()
    add_inline(intro, "아래 자료는 구조 선택의 근거다. 본 P1 데이터에서의 효과나 공식 +3점을 직접 입증하지 않는다.")
    sources = [
        ("MS-TCN (CVPR 2019)", "https://openaccess.thecvf.com/content_CVPR_2019/html/Abu_Farha_MS-TCN_Multi-Stage_Temporal_Convolutional_Network_for_Action_Segmentation_CVPR_2019_paper.html"),
        ("MS-TCN++ 논문", "https://arxiv.org/abs/2006.09220"),
        ("MS-TCN++ 공식 구현", "https://github.com/sj-li/MS-TCN2"),
        ("ASRF 논문", "https://arxiv.org/abs/2007.06866"),
        ("시계열 검증 절차 연구", "https://arxiv.org/abs/1905.11744"),
        ("Benchmark variance 연구", "https://proceedings.mlsys.org/paper_files/paper/2021/hash/0184b0cd3cfb185989f858a1d9f5c1eb-Abstract.html"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(5)
        add_hyperlink(p, label, url)

    p = doc.add_paragraph(style="Heading 2")
    add_inline(p, "재현성 핵심 해시", size=13, color=BLUE)
    rows = [
        ["산출물", "SHA-256"],
        ["terminal_result.json", "7640cc0e29f364a26cd8199a7e9a55acdf329699cd5923679d8f0d513c4af2b1"],
        ["confirmatory_metrics.json", "964cae7d7dbb9f413244462eb9258e883e14f6073ad5549fadf482cb9cd03bd4"],
        ["execution seal", "2d42ce76966876f33daf0bd3e8e62051876f95f92e866588713bcfb84886bb25"],
        ["selected recipe", "171618200c69dc8e5039e5404bdeb4e7cb6369f15ab8ff43af1be7492837515a"],
    ]
    add_data_table(doc, rows)


def build():
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    doc.core_properties.title = "P1 incumbent-preserving MS-TCN++/ASRF v2 연구보고서"
    doc.core_properties.subject = "고용량 구조의 사전등록 실행, 독립 QA, NO_GO 판정"
    doc.core_properties.author = "분당독고다이"
    doc.core_properties.keywords = "P1, MS-TCN++, ASRF, time series, F1, preregistration, NO_GO"
    configure_styles(doc)
    configure_section(doc.sections[0])
    numbering_manager = ensure_numbering(doc)
    add_masthead(doc)
    render_markdown(doc, source_text, numbering_manager)
    add_sources(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
