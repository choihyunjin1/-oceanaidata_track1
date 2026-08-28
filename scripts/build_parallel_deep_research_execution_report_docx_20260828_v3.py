"""Build the verified Korean P1/P2/P3 parallel research decision brief."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = (
    ROOT
    / "reports"
    / "parallel_deep_research_execution_20260828_v3"
    / "P1_P2_P3_병렬_딥리서치_실행보고_20260828.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "5F6B7A"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF4EA"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=9.2, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], LIGHT)
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    mark_header_row(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, *, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    # A one-row callout is still represented as a Word table. Mark that row as
    # its header so screen readers do not report an unlabelled table structure.
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(f"{title}  ")
    set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
    run = paragraph.add_run(body)
    set_run_font(run, size=10.2, color=INK)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead:
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=10.5, bold=True, color=INK)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("분당독고다이 | 병렬 딥리서치 실행 보고"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("2026-08-28  |  "), size=9, color=MUTED)
    add_page_field(footer)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    set_run_font(kicker.add_run("TECHNICAL DECISION BRIEF"), size=9.5, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("P1·P2·P3 병렬 딥리서치 실행 보고"), size=23, bold=True, color="000000")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("구조 탐색, bounded local 실행, 독립 QA, 다음 승격 후보"), size=12.5, color=MUTED)
    for label, value in (
        ("팀", "분당독고다이"),
        ("작성일", "2026-08-28 KST"),
        ("판정 범위", "로컬 연구 전용; 공식 test/sample/submission 접근 및 업로드 없음"),
        ("상태", "세 문제 모두 현재 공식 제출 보류"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}: "), size=10.5, bold=True, color="000000")
        set_run_font(paragraph.add_run(value), size=10.5, color="000000")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(9)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)


def build(output: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_block(doc)
    add_callout(
        doc,
        "최종 결론",
        "이번 사이클에서 즉시 공식 제출할 후보는 없다. 대신 P1은 support, P2는 보정 활성률, P3는 source-transfer와 incumbent 간 간극으로 병목이 분해됐다.",
        fill=PALE_RED,
    )
    add_table(
        doc,
        ["문제", "이번 실험", "핵심 결과", "공식 판정"],
        [
            ["P1", "조건부 실제-event donor", "Support PASS, calibration FAIL", "NO_GO"],
            ["P2", "α40 + quasi-periodic residual", "Δ -0.0000086°C, 1/3 folds", "NO_GO"],
            ["P3", "ERA5 context transfer", "Viewpoint PASS, incumbent FAIL", "NO_GO"],
        ],
        [920, 2600, 3860, 1980],
    )
    add_body(
        doc,
        "공식 33점 스케일과 로컬 물리 단위(°C, m)는 같은 숫자로 비교하지 않는다. 이미 노출된 historical surface는 fresh holdout으로 부르지 않으며, 결과를 보기 전에 구조·분할·gate·prediction hash를 고정했다.",
        bold_lead="판정 원칙. ",
    )

    doc.add_heading("1. P1 — support는 해결, proposal 품질이 새 병목", level=1)
    add_body(
        doc,
        "pre-Q2 실제 장기 이벤트를 동일 station/layer/quarter 정상 구간에 제한적으로 이식했다. CAROTS의 관계보존/교란 증강과 TimeMIL의 event-level localization을 설계 근거로 사용했지만, P1 성능을 보장한다고 해석하지 않았다."
    )
    add_table(
        doc,
        ["지표", "관측", "기준", "판정"],
        [
            ["Train donor support", "22 events / 7 cells", "≥10 / ≥3", "PASS"],
            ["Calibration truth", "26 events / 9 cells", "≥4 / ≥2", "PASS"],
            ["Event recall", "10/26 = 38.46%", "≥80%", "FAIL"],
            ["Row precision", "22.27%", "≥45%", "FAIL"],
            ["Q2 truth access", "0 rows", "0 before gate", "PASS"],
        ],
        [2540, 2380, 2200, 2240],
    )
    add_callout(
        doc,
        "P1 판정",
        "NO_GO_CALIBRATION. 실제 donor가 지원성은 만들었지만 event localization과 정상구간 배제가 부족했다. 다음 한 번은 현재 proposal에 event-level ranker 또는 normality veto를 붙이는 실험이다.",
    )

    doc.add_heading("2. P2 — 안전하지만 99.893%가 no-op", level=1)
    add_body(
        doc,
        "α40을 유지한 채 quasi-periodic residual과 posterior support gate를 적용했다. 3개 architecture-fresh committed forward fold에서 prediction hash를 truth metric 전에 고정했고, out-of-support는 byte-exact α40으로 복귀했다."
    )
    add_table(
        doc,
        ["지표", "관측", "승격 기준", "판정"],
        [
            ["ΔRMSE", "-0.000008606°C", "≤ -0.003°C", "FAIL"],
            ["Bootstrap CI90 upper", "-0.000002253°C", "< 0", "PASS"],
            ["개선 folds", "1/3", "≥2/3", "FAIL"],
            ["보정 활성", "75/69,850 (0.107%)", "정보성 필요", "FAIL"],
            ["Correction RMS", "0.006181°C", "≤0.05°C", "PASS"],
        ],
        [2540, 2380, 2200, 2240],
    )
    add_callout(
        doc,
        "P2 판정",
        "FAIL_GATE_STOP_NO_CSV_NO_RESEARCH_LOOP. 같은 GP의 threshold를 결과에 맞춰 완화하지 않는다. 다음은 α40 profile에 1–2개 vertical-displacement mode만 적용하는 물리적으로 다른 보정이다.",
    )
    doc.add_heading("P2 다음 bounded 계약", level=2)
    add_table(
        doc,
        ["고정 요소", "다음 실험 규칙"],
        [
            ["기준 모델", "공식 검증된 α40을 그대로 유지"],
            ["보정 구조", "실제 수심 기울기 × 1–2개 vertical-displacement mode"],
            ["안전장치", "Density stability / endpoint / exact no-op fallback"],
            ["승격", "Δ≤-0.003°C, CI90 upper<0, ≥2/3 folds"],
        ],
        [2500, 6860],
    )
    add_body(
        doc,
        "내부조석 문헌은 density-surface displacement와 two-mode approximation의 물리적 근거를 제공하지만, deep glider/steric-height의 설명분산을 P2 점수 개선량으로 이전하지 않는다.",
        bold_lead="근거의 경계. ",
    )

    doc.add_page_break()
    doc.add_heading("3. P3 — ERA5 transfer 가치는 입증, incumbent 대체는 실패", level=1)
    add_body(
        doc,
        "기존 v1은 CatBoost import 전에 fit 0회로 종료됐다. 새 dependency-recovery ID에서는 scientific contract를 유지하고 ID와 output만 분리했다. Python/CatBoost stack, 286-feature hash, NaN smoke, deepcopy와 init_model continuation을 lock 전에 검증했다."
    )
    add_table(
        doc,
        ["단계/비교", "ΔRMSE", "강건성", "판정"],
        [
            ["Source vs persistence", "-0.140897m", "3/3 years, CI90 < 0", "PASS"],
            ["Transfer vs local-only", "-0.023440m", "3/3 windows, CI90 < 0", "PASS"],
            ["Transfer vs incumbent", "+0.002325m", "1/3 windows", "FAIL"],
            ["S-ORS critical slice", "+0.021700m", "cap 초과", "FAIL"],
            ["Lead 24h", "-0.024416m", "방향성 개선", "기록"],
        ],
        [2740, 2040, 2840, 1740],
    )
    add_callout(
        doc,
        "P3 해석",
        "ERA5는 약한 CatBoost local-only 모델을 유의하게 끌어올렸지만 강한 incumbent를 통째로 대체하지 못했다. domain classifier AUC 0.9999999이므로 direct pooling/pretrain-only 사용은 계속 금지한다.",
        fill=PALE_GREEN,
    )
    doc.add_heading("4. 다음 실행 우선순위", level=1)
    add_table(
        doc,
        ["순위", "문제", "다음 구조", "중단 기준"],
        [
            ["1", "P3", "Incumbent-preserving ERA5 residual/router", "봉인된 새 구조에서 incumbent 미개선"],
            ["2", "P2", "Two-mode vertical displacement correction", "Δ>-0.003°C 또는 <2 folds"],
            ["3", "P1", "Event-level ranker / normality veto", "Precision<45% 또는 event recall<80%"],
            ["4", "P3", "TimeXer exogenous transfer", "Source pilot gate 실패"],
        ],
        [900, 900, 3860, 3700],
    )
    add_body(
        doc,
        "P3 router는 이번 local slice를 보고 세운 hypothesis-exposed 가설이다. 따라서 confirmatory local claim을 할 수 없으며, 별도 sealed architecture 또는 제한된 공식 probe가 필요하다.",
        bold_lead="주의. ",
    )

    doc.add_page_break()
    doc.add_heading("5. 근거 원장", level=1)
    add_table(
        doc,
        ["ID", "1차 출처", "이번 연구에서의 용도"],
        [
            ["P1-1", "Kim et al., CAROTS, ICML 2025 (PMLR)", "관계보존/교란 증강"],
            ["P1-2", "Chen et al., TimeMIL, ICML 2024 (PMLR)", "Sparse temporal localization"],
            ["P2-1", "Chen et al., Frontiers in Marine Science 2023", "Joint T/S GPR와 uncertainty"],
            ["P2-2", "Bendinger et al., Ocean Science 2024", "Internal-tide two-mode displacement"],
            ["P3-1", "ECMWF/Copernicus ERA5 documentation", "31km atmosphere / 0.36° wave grid"],
            ["P3-2", "CatBoost official GPU documentation", "GPU 비결정성, CPU 선택"],
            ["P3-3", "Wang et al., TimeXer, NeurIPS 2024", "Exogenous cross-attention 후보"],
        ],
        [950, 4350, 4060],
    )
    add_body(
        doc,
        "상세 URL, 주장별 confidence와 limitation은 같은 폴더의 claim-ledger.json에 있다. P1/P2 문헌의 수치 개선을 해커톤 점수로 직접 이전하지 않았고, P3 source 성공을 local incumbent 승리로 과장하지 않았다."
    )

    doc.add_heading("6. 무결성 및 QA", level=1)
    add_table(
        doc,
        ["산출물", "SHA-256"],
        [
            ["P1 result.json", "b32fc6df07b30d315d1d3b09add4455686660ea69ba3b10134ac7f4e0a8c58f4"],
            ["P2 result.json", "c04755750357b8613f7372f98840ba8f8df365173af46524b4be339ee362da2e"],
            ["P3 result.json", "ac92a530d230ea29c475e0b03acb7e16d577633b64401632cebb46fa4e0bbd2f"],
            ["P3 blind seal", "25accc81915e95bebcf4e69cd313b73520c36969b88521a186f5be214c4ba2a7"],
        ],
        [2380, 6980],
    )
    add_callout(
        doc,
        "독립 QA",
        "16/16 cross-problem checks PASS, 관련 pytest 42건 PASS. P3 blind seal은 1,086행이며 target_hs column이 없고, 세 문제 모두 후보 CSV 생성·공식 업로드가 없다.",
        fill=PALE_GREEN,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
