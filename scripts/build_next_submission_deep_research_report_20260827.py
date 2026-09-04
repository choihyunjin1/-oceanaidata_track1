"""Build the Korean next-submission deep-research decision memo."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(r"C:\Users\cedis\PycharmProjects\PythonProject")
OUTPUT = (
    REPO
    / "reports"
    / "next_submission_deep_research_20260827_v1"
    / "다음_제출기회_고가치_딥리서치_20260827.docx"
)

NAVY = "12304A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "287C78"
MUTED = "5A6573"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8F0F7"
PALE_GREEN = "E7F2EC"
PALE_GOLD = "FFF3D6"
PALE_RED = "FBE9E7"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111111"


def set_run_font(run, size=10.5, color=BLACK, bold=None, italic=None):
    run.font.name = "Arial"
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=105, bottom=70, end=105):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths, header=True, header_fill=LIGHT_GRAY, font_size=8.8):
    assert sum(widths) == 9360
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout")) or OxmlElement("w:tblLayout")
    if layout.getparent() is None:
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(1.5)
                para.paragraph_format.line_spacing = 1.0
                for run in para.runs:
                    set_run_font(run, font_size, bold=(header and row_idx == 0))
        if header and row_idx == 0:
            for cell in row.cells:
                set_fill(cell, header_fill)
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = str(text)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
    format_table(table, widths, header_fill=header_fill, font_size=font_size)
    return table


def add_para(doc, text="", *, bold=False, color=BLACK, size=10.5, after=5, italic=False, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.08
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_run_font(run, size, color, bold, italic)
    return para


def add_mixed(doc, segments, *, after=4, size=10.5):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.08
    for text, bold, color in segments:
        run = para.add_run(text)
        set_run_font(run, size, color, bold)
    return para


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_source(doc, text):
    para = doc.add_paragraph(style="Table Source")
    para.add_run(text)
    return para


def add_callout(doc, title, body, fill, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    set_fill(cell, fill)
    set_cell_margins(cell, 120, 150, 120, 150)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    set_run_font(p1.add_run(title), 11.3, title_color, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    set_run_font(p2.add_run(body), 10.2)
    format_table(table, [9360], header=False, font_size=10.2)
    add_para(doc, "", after=1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.3, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in (
        ("Heading 1", 15.5, BLUE, 11, 5),
        ("Heading 2", 12.5, DARK_BLUE, 8, 4),
        ("Heading 3", 11.2, TEAL, 6, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    source = doc.styles.add_style("Table Source", 1)
    source.font.name = "Arial"
    source.font.size = Pt(8.1)
    source.font.italic = True
    source.font.color.rgb = RGBColor.from_string(MUTED)
    source._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    source._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    source._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    source.paragraph_format.space_before = Pt(3)
    source.paragraph_format.space_after = Pt(4)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.82)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.38)
        section.footer_distance = Inches(0.38)
        hp = section.header.paragraphs[0]
        hp.paragraph_format.space_after = Pt(0)
        set_run_font(hp.add_run("분당독고다이  |  다음 제출기회 의사결정 메모"), 8.2, MUTED, True)
        add_page_number(section.footer.paragraphs[0])


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    add_para(doc, "DECISION MEMO", bold=True, color=NAVY, size=22, after=2)
    add_para(doc, "다음 제출기회를 값지게 쓰는 P2·P3 딥리서치", bold=True, color=NAVY, size=15.2, after=9)
    for label, value in (
        ("팀", "분당독고다이"),
        ("기준", "2026-08-27 KST"),
        ("상태", "문헌·계보 추적 및 신규 로컬 반증 실험 완료 / 공식 업로드 미수행"),
        ("결정", "P2 구조 probe 1장 우선, P3 같은 축 추가 미세조정 보류"),
    ):
        add_mixed(doc, [(f"{label}: ", True, BLACK), (value, False, BLACK)], after=2)
    add_para(doc, "", after=2)

    add_callout(
        doc,
        "결론부터",
        "다음 한 장의 1순위는 P2-SEASONAL-OAS-TS-10-PROJECTED입니다. 현재 P2 최고본 U를 90% 보존하고, 같은 시각 공개층의 수직 T/S 공분산 복원을 10%만 섞은 뒤 기존 물리 투영을 적용합니다. 로컬 총 RMSE는 −0.007782℃, 투영 후 −0.008060℃ 개선됐지만 한 블록은 악화했으므로 ‘배포 확정’이 아니라 사전 등록된 공식 구조 probe입니다.",
        PALE_GREEN,
    )
    add_table(
        doc,
        ["문제", "현재 공식 최고", "관측 headroom", "다음 슬롯"],
        [
            ["P2", "26.611283점 / 0.535727℃", "약 +1.991점", "OAS 구조 probe 1순위"],
            ["P3", "24.066167점 / 0.583892m", "약 +0.099점", "ERA5 또는 새 구조까지 보류"],
        ],
        [1200, 2900, 1800, 3460],
        header_fill=PALE_BLUE,
    )
    add_source(doc, "공식 상태: finite_horizon_submission_decision_20260827_v1 receipt. P3는 공개축 이차최적점에서 이미 +0.240938점 개선.")

    add_heading(doc, "1. 왜 P2인가", 1)
    add_para(doc, "P2는 61일 동안 2·3·4층의 T/S를 복원하지만 같은 시각의 1·5·6·7층 T/S는 공개된다. 기존 최고 계보는 긴 시계열·다중 deep expert와 층별 공식축 후처리를 이미 사용한다. 이전 full-mask BiTCN, CSDI/SSSD형, 연주기 anomaly transfer, 조석/RTS residual이 탈락했으므로 더 긴 창이나 더 큰 모델만 추가하는 것은 새로운 가설이 아니다.")
    add_callout(
        doc,
        "새 오류공간",
        "시간 방향 외삽이 아니라 같은 시각 수직 프로파일의 T/S 교차 공분산을 직접 조건부 복원합니다. 기존 모델이 놓친 작은 수직 구조만 10% 더하므로 현재 최고본을 크게 훼손하지 않습니다.",
        PALE_BLUE,
    )
    add_para(doc, "Conditional multivariate FPCA 연구는 부분 관측된 다변량 프로파일에서 교차 공분산으로 잠재 score를 조건부 추정하는 원리를 제시한다. 우리 실험은 그 효과 크기를 전이하지 않고, 동일한 조건부 선형 복원 원리만 OAS shrinkage로 구현했다.")
    add_source(doc, "Primary source: https://arxiv.org/html/2608.05376v1 (2026 arXiv preprint; domain transfer caveat applies).")

    add_heading(doc, "2. 반증을 먼저 통과시켰다", 1)
    add_heading(doc, "2.1 경계 정합 전년 프로파일 — 기각", 2)
    add_table(
        doc,
        ["노출 블록", "reference", "후보", "10% 혼합", "판정"],
        [
            ["2025 May–Jun", "1.286492", "1.824135", "1.298818", "악화"],
            ["2025 Jul–Aug", "1.103796", "4.924675", "1.274677", "큰 악화"],
        ],
        [2350, 1700, 1700, 1700, 1910],
        header_fill=PALE_RED,
    )
    add_source(doc, "artifacts/p2_boundary_registered_prior_20260827_v1/result.json. 두 블록의 oracle α도 음수.")
    add_para(doc, "이전 해의 같은 계절 프로파일을 공개층과 양쪽 7일 경계에 등록해도 계절 상태 전이와 깊은 층 편향을 견디지 못했다. 공식 제출 후보에서 제외한다.")

    add_heading(doc, "2.2 계절 국소 OAS 프로파일 — 제한적 통과", 2)
    add_table(
        doc,
        ["노출 블록", "reference", "OAS 단독", "10% 혼합", "ΔRMSE"],
        [
            ["2024 Sep–Oct", "0.447793", "0.746001", "0.433021", "−0.014772"],
            ["2025 Jul–Aug", "1.053477", "1.341220", "1.066635", "+0.013157"],
            ["2025 Nov–Dec", "0.613081", "0.252229", "0.550264", "−0.062817"],
            ["전체", "0.768367", "—", "0.760586", "−0.007782"],
        ],
        [2300, 1650, 1650, 1700, 2060],
        header_fill=PALE_GREEN,
    )
    add_source(doc, "artifacts/p2_oas_conditional_profile_20260827_v3/result.json; 총 69,850행.")
    add_table(
        doc,
        ["추가 증거", "결과", "해석"],
        [
            ["물리 투영 후", "0.760308 / Δ −0.008060℃", "투영 11,062행(15.84%) 작동"],
            ["KST-day bootstrap", "90% CI [−0.012785, −0.003176]", "5,000회; 개선 확률 99.92%"],
            ["같은 계절 2024 Sep–Oct", "9주·L2/L3/L4 모두 개선", "공식 gap과 계절 정합"],
            ["반증", "2025 Jul–Aug 악화", "regime shift 위험 보존"],
        ],
        [2700, 3000, 3660],
        header_fill=PALE_GOLD,
    )

    add_heading(doc, "3. 사전 등록 후보", 1)
    add_callout(
        doc,
        "P2-SEASONAL-OAS-TS-10-PROJECTED",
        "최종 예측 = ProfileProjection(0.90 × current P2 U + 0.10 × seasonal local OAS conditional profile). α=0.10은 고정하며 공식 결과를 본 뒤 같은 날 재탐색하지 않습니다.",
        PALE_GREEN,
    )
    add_table(
        doc,
        ["항목", "고정 사양"],
        [
            ["season partition", "14일 bin, 원형 day-of-year ±60일"],
            ["공개 입력 X", "T/S 1·5·6·7층 + 연주기 sin/cos 1~4차"],
            ["출력 Y", "T/S 2·3·4층 joint profile"],
            ["추정", "OAS shrinkage covariance, 행별 관측 X 부분집합 조건부 평균"],
            ["누수 금지", "공식 gap의 은닉층 label은 적합에 사용하지 않음"],
            ["후처리", "기존 endpoint/PAVA profile projection 정확히 1회"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "4. 제출 전 승격 게이트", 1)
    add_table(
        doc,
        ["단계", "통과 조건", "실패 시"],
        [
            ["P0 무결성", "26,061 exact keys/order; finite/range; gap label 미사용; hash·환경 receipt; 재실행 일치", "NO-GO"],
            ["P1 로컬", "총 projected Δ≤−0.0075℃ 재현; 같은 계절 3층·9주 개선; bootstrap 상한<0; U 대비 차이 RMS≥0.02℃", "NO-GO 또는 원인 감사"],
            ["P2 공식", "한 장만 제출하고 사전 결정표대로 판독", "같은 축 당일 추가 sweep 금지"],
        ],
        [1500, 5900, 1960],
        header_fill=PALE_BLUE,
        font_size=8.5,
    )
    add_callout(
        doc,
        "승격 상태",
        "P0·P1을 통과해도 OFFICIAL_PROBE_ELIGIBLE입니다. 세 외부 블록이 이미 노출됐고 한 블록이 악화했으므로 DEPLOYMENT_GO 또는 +3점 보장을 주장하지 않습니다.",
        PALE_GOLD,
        title_color=RED,
    )

    add_heading(doc, "5. 공식 한 장의 사전 결정표", 1)
    add_table(
        doc,
        ["공식 결과", "조치", "해석"],
        [
            [">26.611283점, 반올림 잡음 초과", "새 P2 public best 채택", "수직 covariance expert 유효; α 즉시 재튜닝 금지"],
            ["동률·경계", "방향성 기록, 당일 중단", "정보 부족; 같은 축 반복 금지"],
            ["악화", "OAS 축 종료", "남은 슬롯은 다른 family에 보존"],
        ],
        [2650, 2800, 3910],
        header_fill=PALE_GREEN,
    )
    add_para(doc, "로컬 RMSE Δ를 공식 점수 Δ로 직접 환산하지 않는다. 과거 로컬-공식 민감도 괴리는 이 후보를 무조건 탈락시키지 않을 이유이지, 점수 상승을 보장할 이유가 아니다.")

    add_heading(doc, "6. 후보별 기회가치", 1)
    add_table(
        doc,
        ["후보", "구조적 새로움", "실패 시 정보", "권고"],
        [
            ["P2 OAS 10% + projection", "높음", "높음: profile covariance 가설 판별", "1순위"],
            ["P2 경계 정합 전년 prior", "중간", "이미 반증 완료", "제출 금지"],
            ["P3 −8/−12", "낮음: 같은 1D축", "낮음", "보류"],
            ["generic 대형 imputer", "중간", "준비·QA 비용 큼", "다음 연구"],
            ["P3 ERA5 context transfer", "높음", "높음", "고정 실험 완료 후"],
        ],
        [2900, 2000, 2800, 1660],
    )
    add_source(doc, "P2 관측 headroom은 P3보다 약 20배 크다. P3 same-axis optimum은 이미 공식 확인됨.")

    add_heading(doc, "7. 다음 큰 구조", 1)
    add_para(doc, "OAS가 공식 통과하면 이를 고정 수직 expert로 삼아 LSTI형 장·단기 양방향 imputer의 gating 입력으로 확장한다. ImputeFormer의 저랭크 attention은 긴 block missing의 계산 효율 후보지만, 교통 센서 benchmark를 그대로 전이하지 않는다. CSDI/SSSD형은 로컬에서 이미 weight 0이므로 단순 재실행하지 않는다. 모든 다음 딥 모델은 정확한 61일 contiguous mask와 공개층 조건을 유지한다.")
    add_table(
        doc,
        ["근거", "핵심", "우리 판단"],
        [
            ["LSTI / TMLR", "장·단기 양방향 expert + meta weighting", "OAS 공식 통과 뒤 2단계"],
            ["ImputeFormer / KDD 2024", "저랭크 projected attention", "효율 후보, 도메인 재검증 필요"],
            ["CSDI / NeurIPS 2021", "조건부 diffusion", "로컬 유사 family 실패로 후순위"],
            ["TSI-Bench", "다중 missing pattern 대규모 비교", "61일 block 계약 고정 근거"],
        ],
        [2600, 3350, 3410],
        header_fill=PALE_BLUE,
    )

    doc.add_page_break()
    add_heading(doc, "8. 한계·중단 조건", 1)
    for item in (
        "세 외부 블록은 fresh holdout이 아니며, bootstrap은 적응적 모델 선택 편향을 제거하지 못합니다.",
        "frozen reference는 공식 lineage OOF이며 현재 U의 미세 공식축 보정을 완전히 재현한 것은 아닙니다.",
        "2025 Jul–Aug 악화는 계절 covariance의 regime risk를 보여줍니다.",
        "정확한 대회명·파일명·팀명으로 공개 참가자 코드를 검색했지만 신뢰 가능한 저장소를 찾지 못했습니다. 이는 부재의 증명이 아닙니다.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            set_run_font(run, 9.8)
        if not p.runs:
            set_run_font(p.add_run(item), 9.8)
        else:
            p.runs[0].text = item
            set_run_font(p.runs[0], 9.8)
    add_callout(
        doc,
        "최종 권고",
        "P0/P1을 먼저 재현한 뒤 P2 한 장만 사용하십시오. 성공하면 새 수직 expert를 다음 backbone에 편입하고, 동률·실패면 같은 축을 닫습니다. P3는 ERA5 또는 완전히 다른 구조가 준비될 때까지 보존합니다. 최소 +3 공식점은 장기 목표로 유지하되, 이번 한 장은 그 목표에 필요한 구조 가설을 공식적으로 판별하는 데 씁니다.",
        PALE_GREEN,
    )

    add_heading(doc, "근거 출처·재현 경로", 1)
    sources = (
        "Conditional multivariate FPCA: https://arxiv.org/html/2608.05376v1",
        "ImputeFormer: https://arxiv.org/html/2312.01728v3 · https://github.com/tongnie/ImputeFormer",
        "LSTI: https://openreview.net/forum?id=9NVJ0ZgEfT",
        "CSDI: https://proceedings.neurips.cc/paper/2021/hash/cfe8504bda37b575c70ee1a8276f3486-Abstract.html",
        "TSI-Bench: https://arxiv.org/abs/2406.12747 · https://github.com/WenjieDu/Awesome_Imputation",
        "Leaderboard risk: https://proceedings.mlr.press/v37/blum15.html · https://www.jmlr.org/papers/v11/cawley10a.html",
        "Local result: artifacts/p2_oas_conditional_profile_20260827_v3/result.json",
        "Full report and claim ledger: reports/next_submission_deep_research_20260827_v1/",
    )
    for item in sources:
        add_para(doc, item, color=MUTED, size=8.6, after=2)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
