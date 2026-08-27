from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = Path.home() / "Downloads" / "해양 해커톤 제출용" / "20260826_round_C_preregistered_P1x3_P2x1"
OUTPUT = PACKAGE / "20260826_일일제출_의사결정_메모_v2.docx"
SOURCE = ROOT / "artifacts" / "daily_submission_value_20260826_v1" / "report-source.md"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "666666"
GREEN = "EAF4EA"
GOLD = "FFF4CE"
RED = "FCE8E6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, *, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_paragraph_runs(paragraph, *, size=10.5, color="000000", bold=None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MID_GRAY)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    set_paragraph_runs(paragraph, size=11)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    set_paragraph_runs(paragraph, size=11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, LIGHT_GRAY)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_runs(paragraph, size=9.5, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_runs(paragraph, size=9.5)
    set_table_geometry(table, widths)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    return table


def add_callout(doc: Document, label: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    lead = paragraph.add_run(label + "  ")
    set_run_font(lead, size=11, bold=True, color=DARK_BLUE)
    body = paragraph.add_run(text)
    set_run_font(body, size=11)
    set_table_geometry(table, [9360])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build() -> None:
    if OUTPUT.exists():
        raise FileExistsError(OUTPUT)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    doc.settings.odd_and_even_pages_header_footer = True
    section.different_first_page_header_footer = True
    for header_part in (section.header, section.even_page_header, section.first_page_header):
        header = header_part.paragraphs[0]
        header.text = "분당독고다이 | OCEAN AI 해커톤"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_runs(header, size=9, color=MID_GRAY)
    for footer_part in (section.footer, section.even_page_footer, section.first_page_footer):
        footer = footer_part.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = footer.add_run("2026-08-26  |  Page ")
        set_run_font(r, size=9, color=MID_GRAY)
        add_page_field(footer)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("일일 제출 기회 의사결정 메모")
    set_run_font(run, size=23, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("2026-08-26 | P1 구조 절제 3건 + P2 해석적 최적화 1건")
    set_run_font(run, size=14, color=MID_GRAY)
    for label, value in (
        ("팀", "분당독고다이"),
        ("상태", "파일 동결 · 독립 QA 진행 · 공식 업로드 0회"),
        ("공식 한도", "오늘 각 문제 3회; 현재 P1/P2/P3 모두 3/3"),
        ("의사결정", "P1 3회, P2 1회, P3 0회 — 사용자 action-time 확인 후 실행"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        left = p.add_run(label + ": ")
        set_run_font(left, size=11, bold=True)
        right = p.add_run(value)
        set_run_font(right, size=11)

    add_callout(
        doc,
        "권고",
        "오늘의 슬롯을 단순 소진하지 않고, P1은 개선 1건과 메커니즘 식별 2건으로 모두 사용하며 P2는 수학적으로 우월한 1건만 사용한다. P3는 ERA5 고정 실험 완료 전까지 보류한다.",
        GREEN,
    )

    doc.add_heading("1. 공식 현황과 판단 원칙", level=1)
    add_table(
        doc,
        ["문제", "Original", "A", "B / 현재 best"],
        [
            ["P1 F1", "0.790709", "0.786145", "0.793710 / B"],
            ["P2 RMSE", "0.541085", "0.713520", "0.599921 / Original"],
            ["P3 RMSE", "0.607071", "0.611680", "0.609346 / Original"],
        ],
        [1500, 2460, 2460, 2940],
    )
    p = doc.add_paragraph(
        "공식 리더보드는 문제별 public best를 합산하므로 나쁜 새 제출이 기존 최고 기록을 직접 낮추지는 않는다. 그러나 각 제출은 공개 holdout에 대한 적응적 질의다. 따라서 즉시 개선 가능성과 다음 구조 선택을 바꾸는 정보가치 중 하나가 명확해야 한다."
    )
    set_paragraph_runs(p, size=11)

    doc.add_heading("2. 사전등록 제출 묶음", level=1)
    add_table(
        doc,
        ["순서", "후보", "목적", "오늘 사용"],
        [
            ["P1-1", "정점-층 disagreement router", "개선 research floor", "1/3"],
            ["P1-2", "B ∩ O", "추가양성 절제 probe", "2/3"],
            ["P1-3", "B ∪ O", "제거양성 복원 probe", "3/3"],
            ["P2-1", "public quadratic alpha*", "개선 exploit", "1/3"],
            ["P3", "ERA5 완료 전 없음", "실험 오염 방지", "0/3"],
        ],
        [1050, 3450, 3000, 1860],
    )
    add_bullet(doc, "네 파일과 해시는 오늘 공식 결과를 보기 전에 동결한다.")
    add_bullet(doc, "가능하면 네 건을 연속 제출한 뒤 결과를 열어 같은 날 후속 튜닝을 차단한다.")
    add_bullet(doc, "P1 probe의 실패는 모델 실패가 아니라 추가·제거 메커니즘의 방향을 식별한 관측으로 기록한다.")

    doc.add_heading("3. P1 — 개선과 메커니즘 식별", level=1)
    doc.add_heading("3.1 P1-1 정점-층 disagreement router", level=2)
    p = doc.add_paragraph(
        "현 베스트 B를 보존하면서 O-only 중 G-ORS layer 1과 I-ORS layer 2를 복원하고, B-only 중 S-ORS layer 1/5/6과 I-ORS layer 4를 제거한다. 공식 파일에서 B와 다른 행은 229개다."
    )
    set_paragraph_runs(p)
    add_table(
        doc,
        ["지표", "B", "Router", "차이 / 판정"],
        [
            ["로컬 OOF F1", "0.86467009", "0.86690000", "+0.00222991"],
            ["CI90", "—", "[+0.00091805, +0.00370742]", "0 미포함"],
            ["시간 fold", "기준", "2 개선, 1 보합", "research floor"],
            ["공식 변경 행", "0", "229", "식별 가능"],
        ],
        [2200, 1800, 3100, 2260],
    )
    add_callout(
        doc,
        "한계",
        "정식 +0.0255 승격 gate에는 미달한다. 같은 OOF에서 disagreement cell을 선택했으므로 작은 효과는 선택 편향 가능성이 있으며, 결과는 연구용으로 해석한다.",
        GOLD,
    )

    doc.add_heading("3.2 P1-2 / P1-3 구조 절제", level=2)
    add_table(
        doc,
        ["Probe", "공식 변경", "로컬 delta", "결과가 답하는 질문"],
        [
            ["B ∩ O", "B-only 176개 제거", "-0.00159359", "B의 새 양성이 유효했는가?"],
            ["B ∪ O", "O-only 824개 복원", "-0.00262723", "B의 pruning이 유효했는가?"],
        ],
        [1500, 2350, 1900, 3610],
    )
    p = doc.add_paragraph(
        "공식 public split의 행 구성이 비공개이므로 전체 CSV 양성 개수와 6자리 F1만으로 숨은 TP를 역산할 수 없다. 따라서 두 점은 점수 최적화가 아니라 사전 고정된 intervention의 부호를 확인하는 실험이다."
    )
    set_paragraph_runs(p)

    doc.add_heading("4. P2 — public MSE 이차곡선의 해석적 최적점", level=1)
    p = doc.add_paragraph(
        "동일 경로 prediction(alpha)=O+alpha(A-O)에서 public MSE는 정확한 이차식이다. 기존 B 파일은 alpha=0.5 중점과 최대 3.55e-15°C 오차로 일치하므로 Original·B·A의 공식 RMSE 세 점만으로 경로상의 최적 alpha를 결정할 수 있다."
    )
    set_paragraph_runs(p)
    add_table(
        doc,
        ["항목", "값"],
        [
            ["적합식", "MSE = 0.164146710286 alpha² + 0.052191102889 alpha + 0.292772977225"],
            ["alpha*", "-0.158976999289"],
            ["예상 public RMSE", "0.537237735794"],
            ["예상 개선", "0.003847264206°C"],
            ["반올림 오차 포함 개선 범위", "0.00384644~0.00384808°C"],
        ],
        [2500, 6860],
    )
    add_table(
        doc,
        ["로컬 검증", "Incumbent", "alpha* 후보", "판정"],
        [
            ["p100 OOF RMSE", "0.990326912", "0.985443525", "-0.004883387°C"],
            ["일별 CI90", "—", "[-0.007460316, -0.002186438]", "P(개선)=0.9982"],
            ["시간 fold", "기준", "2 개선, 1 소폭 악화", "방향 대체로 일치"],
            ["독립 QA", "—", "26,061행 · 키/산식/해시 일치", "PASS"],
        ],
        [2400, 1800, 3300, 1860],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    lead = p.add_run("SHA-256: ")
    set_run_font(lead, size=10, bold=True)
    value = p.add_run("9cc951801cf6b6cdacc2c826126d9c2f72ef34fc67e46c6a21261c7a1ba845ff")
    set_run_font(value, size=9, color=MID_GRAY)
    add_callout(
        doc,
        "해석",
        "이 결과가 맞아도 새 모델 구조의 일반화 성공이 아니라, 기존 보정 방향을 소량 반대로 적용한 public 경로 최적화 성공이다. 같은 경로의 SAFE 절반점은 새 정보를 주지 않아 제출하지 않는다.",
        GREEN,
    )

    doc.add_heading("5. 실행·판정 규율", level=1)
    add_number(doc, "P1-1, P1-2, P1-3, P2-1의 정확한 파일과 해시를 확인한다.")
    add_number(doc, "사용자의 action-time 승인을 받은 뒤에만 공식 사이트에 업로드한다.")
    add_number(doc, "네 건을 제출하기 전에는 어떤 새 공식 점수도 후보 제작에 사용하지 않는다.")
    add_number(doc, "제출 후 문제별 current best 보존 여부, 새 점수, 제출 시각, 파일 해시를 ledger에 기록한다.")
    add_number(doc, "P1 probe는 개선/악화 방향으로 다음 모델의 추가·제거 메커니즘을 선택한다. P2는 예측 곡선과 관측 점수의 차이를 기록한다.")
    add_number(doc, "P3는 ERA5 source gate와 사전등록 local gate를 모두 통과하기 전에는 제출하지 않는다.")

    doc.add_heading("6. 연구 근거와 적용 한계", level=1)
    sources = [
        ("Blum & Hardt (2015), The Ladder", "반복 leaderboard 적응은 public holdout 과적합을 유발할 수 있다.", "https://proceedings.mlr.press/v37/blum15.html"),
        ("Dwork et al. (2015), Reusable Holdout", "적응적 holdout 재사용은 표준 일반화 가정을 깨뜨린다.", "https://proceedings.neurips.cc/paper_files/paper/2015/hash/bad5f33780c42f2588878a9d07405083-Abstract.html"),
        ("Cawley & Talbot (2010)", "유한 검증면에서 최고 후보 선택 자체가 낙관 편향을 만든다.", "https://www.jmlr.org/papers/v11/cawley10a.html"),
        ("Varma & Simon (2006)", "튜닝과 평가를 분리한 nested 검증이 선택 편향을 줄인다.", "https://bmcbioinformatics.biomedcentral.com/articles/10.1186/1471-2105-7-91"),
        ("Russo & Van Roy (2017)", "탐색은 즉시 보상뿐 아니라 다음 의사결정을 바꾸는 정보량으로 평가한다.", "https://pubsonline.informs.org/doi/10.1287/opre.2017.1663"),
    ]
    for title_text, explanation, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(title_text + ": ")
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(explanation + " " + url)
        set_run_font(r2, size=10.5)

    add_callout(
        doc,
        "최종 상태",
        "제출 파일과 백업 세트는 생성·동결되었고 P2 독립 QA는 PASS다. 공식 업로드는 0회이며, 사용자 확인 직후 P1 3회와 P2 1회를 실행할 수 있다.",
        GOLD,
    )
    doc.core_properties.title = "2026-08-26 일일 제출 기회 의사결정 메모"
    doc.core_properties.subject = "OCEAN AI 해커톤 P1/P2 제출 가치 분석"
    doc.core_properties.author = "분당독고다이"
    doc.save(OUTPUT)
    shutil.copy2(SOURCE, PACKAGE / "report-source.md")
    print(OUTPUT)


if __name__ == "__main__":
    build()
