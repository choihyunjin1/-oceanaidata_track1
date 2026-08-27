from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


HERE = Path(__file__).resolve().parent
DOCX = HERE / "ocean_hackathon_breakthrough_deep_research_20260826_ko.docx"
OUT = HERE / "structural_qa.json"
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "dc": "http://purl.org/dc/elements/1.1/",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
}


def main():
    with zipfile.ZipFile(DOCX) as archive:
        bad_member = archive.testzip()
        xml = ET.fromstring(archive.read("word/document.xml"))
        settings = ET.fromstring(archive.read("word/settings.xml"))
        numbering = ET.fromstring(archive.read("word/numbering.xml"))
        footer = ET.fromstring(archive.read("word/footer1.xml"))
        core = ET.fromstring(archive.read("docProps/core.xml"))

    doc = Document(DOCX)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    all_text += "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    headings = {}
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            headings[p.style.name] = headings.get(p.style.name, 0) + 1
    num_ids = [
        x.get(f"{{{NS['w']}}}val")
        for x in xml.findall(".//w:numPr/w:numId", NS)
    ]
    author = core.find("dc:creator", NS)
    modified = core.find("cp:lastModifiedBy", NS)
    key_phrases = [
        "P2 comparator de-bias",
        "P1 anchor/postprocess parity",
        "F1-aware non-destructive interval rescue",
        "shrinkage conditional hydrographic FPCA",
        "Pre-cycle audit 이후 정확히 두 학습 사이클",
        "P3 고정 실험 및 공식 test/sample/submission/candidate 미접근",
    ]
    result = {
        "file": str(DOCX),
        "bytes": DOCX.stat().st_size,
        "sha256": hashlib.sha256(DOCX.read_bytes()).hexdigest(),
        "zip_test_bad_member": bad_member,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": headings,
        "section_count": len(doc.sections),
        "section_geometry": [
            {
                "page_width_in": round(s.page_width.inches, 3),
                "page_height_in": round(s.page_height.inches, 3),
                "left_margin_in": round(s.left_margin.inches, 3),
                "right_margin_in": round(s.right_margin.inches, 3),
                "top_margin_in": round(s.top_margin.inches, 3),
                "bottom_margin_in": round(s.bottom_margin.inches, 3),
                "header_in": round(s.header_distance.inches, 3),
                "footer_in": round(s.footer_distance.inches, 3),
            }
            for s in doc.sections
        ],
        "numbered_paragraphs": len(num_ids),
        "distinct_numbering_instances_used": len(set(num_ids)),
        "abstract_numbering_definitions": len(numbering.findall("w:abstractNum", NS)),
        "hyperlinks": len(xml.findall(".//w:hyperlink", NS)),
        "manual_page_breaks": len(xml.findall(".//w:br[@w:type='page']", NS)),
        "page_break_before_count": len(xml.findall(".//w:pageBreakBefore", NS)),
        "update_fields": settings.find("w:updateFields", NS) is not None,
        "footer_page_field": any(
            "PAGE" in (node.text or "") for node in footer.findall(".//w:instrText", NS)
        ),
        "author_blank": author is None or not (author.text or "").strip(),
        "last_modified_by_blank": modified is None or not (modified.text or "").strip(),
        "key_phrases": {phrase: phrase in all_text for phrase in key_phrases},
        "forbidden_tokens_absent": all(
            token not in all_text
            for token in (":codex-file-citation", "turn688view", "PLACEHOLDER", "TODO")
        ),
    }
    result["pass"] = all(
        [
            result["zip_test_bad_member"] is None,
            result["tables"] == 4,
            result["section_count"] == 1,
            result["update_fields"],
            result["footer_page_field"],
            result["author_blank"],
            result["last_modified_by_blank"],
            all(result["key_phrases"].values()),
            result["forbidden_tokens_absent"],
        ]
    )
    OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
