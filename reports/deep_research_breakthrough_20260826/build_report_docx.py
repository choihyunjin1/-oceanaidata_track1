from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "report-source.md"
OUTPUT = HERE / "ocean_hackathon_breakthrough_deep_research_20260826_ko.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "111827"
GREEN = "245B3A"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_east_asia_font(run, latin="Calibri", east_asia="Malgun Gothic"):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def set_run_style(run, *, size=11, bold=None, italic=None, color=BLACK, latin="Calibri"):
    set_east_asia_font(run, latin=latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def shade_paragraph(paragraph, fill, left_border=None):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    if left_border:
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr")
            ppr.append(pbdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        pbdr.append(left)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for side, value in CELL_MARGINS.items():
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "D0D5DD")

    grid = tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440.0)


def mark_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def add_hyperlink(paragraph, text, url, *, size=10, color=BLUE):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rpr.append(fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|https?://\S+)"
)


def add_inline(paragraph, text, *, size=11, color=BLACK, bold=False, table=False):
    text = text.replace("–", "-").replace("—", "-")
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_style(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_style(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_style(run, size=max(8.3, size - 0.5), color=DARK_BLUE, latin="Consolas")
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url, size=size)
        else:
            clean = token.rstrip(".,;)")
            suffix = token[len(clean) :]
            add_hyperlink(paragraph, "원문" if table else clean, clean, size=size)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_style(run, size=size, color=color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_style(run, size=size, color=color, bold=bold)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_numbering(doc, *, bullet):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def attach_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(num)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_style(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        wrapper = OxmlElement("w:r")
        wrapper.append(node)
        paragraph._p.append(wrapper)


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.clear()
    set_paragraph_spacing(hp, after=0, line=1.0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = hp.add_run("DEEP RESEARCH / 분당독고다이")
    set_run_style(left, size=8.5, bold=True, color=MUTED)
    right = hp.add_run("\t2026-08-26")
    set_run_style(right, size=8.5, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.clear()
    set_paragraph_spacing(fp, after=0, line=1.0)
    add_page_field(fp)


def add_title_block(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=3, line=1.0)
    r = p.add_run("MODEL BREAKTHROUGH RESEARCH")
    set_run_style(r, size=9.5, bold=True, color=BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=5, line=1.0)
    r = p.add_run("해양 해커톤 모델 돌파구 딥리서치")
    set_run_style(r, size=25, bold=True, color=NAVY)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14, line=1.15)
    r = p.add_run("코드 취약점, 외부 연구 근거, 2사이클 검증 설계")
    set_run_style(r, size=13, color=MUTED)

    for label, value in [
        ("팀", "분당독고다이"),
        ("작성일", "2026-08-26"),
        ("범위", "P1 이상탐지 / P2 수온복원 / 검증·제출전략"),
        ("보호", "P3 고정 실험 및 공식 test/sample/submission/candidate 미접근"),
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.0)
        r = p.add_run(f"{label}: ")
        set_run_style(r, size=10.2, bold=True, color=DARK_BLUE)
        r = p.add_run(value)
        set_run_style(r, size=10.2, color=BLACK)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4, line=1.0)


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def choose_widths(headers):
    n = len(headers)
    if n == 2:
        return [3900, 5460]
    if n == 3:
        return [1700, 3000, 4660]
    if n == 4:
        return [900, 2520, 2920, 3020]
    if n == 5 and headers[0] == "문제":
        return [850, 1600, 2200, 1450, 3260]
    if n == 5:
        return [700, 2450, 2200, 2850, 1160]
    base = CONTENT_WIDTH_DXA // n
    widths = [base] * n
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def render_claim_ledger(doc, rows):
    for row in rows[1:]:
        if len(row) < 5:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=4, after=2, line=1.08)
        r = p.add_run(f"{row[0]}  {row[1]}")
        set_run_style(r, size=10.2, bold=True, color=DARK_BLUE)
        p.paragraph_format.keep_with_next = True
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=3, line=1.08)
        add_inline(p, f"내부: {row[2]}  |  외부: {row[3]}  |  신뢰도: {row[4]}", size=9.2, color=MUTED)
        shade_paragraph(p, CALLOUT, left_border="B8C4D6")


def render_table(doc, rows):
    if not rows:
        return
    if "내부 근거" in rows[0] and "외부 1차 근거" in rows[0]:
        render_claim_ledger(doc, rows)
        return
    headers = rows[0]
    widths = choose_widths(headers)
    table = doc.add_table(rows=len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    mark_repeat_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.05)
            if c_idx == 0 and len(headers) <= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(
                p,
                value,
                size=8.6 if len(headers) >= 4 else 9.0,
                color=BLACK,
                bold=r_idx == 0,
                table=True,
            )
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
    after = doc.add_paragraph()
    set_paragraph_spacing(after, before=0, after=2, line=1.0)


def add_callout(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=10, line=1.16)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.10)
    add_inline(p, text.replace("[결론]", "결론:"), size=11, color=NAVY)
    shade_paragraph(p, LIGHT_BLUE, left_border=BLUE)


def render_markdown(doc, lines, bullet_num):
    start = 0
    while start < len(lines) and not lines[start].startswith("> "):
        start += 1
    i = start
    page_break_heads = {
        "4. P1 코드 감사: 현재 접근이 놓친 것",
        "6. P2 코드 감사: 표현력보다 수직구조와 수송성",
        "8. Pre-cycle audit 이후 정확히 두 학습 사이클",
    }
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("> "):
            add_callout(doc, line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            p = doc.add_paragraph(style="Heading 1")
            if title in page_break_heads:
                p.paragraph_format.page_break_before = True
            add_inline(p, title, size=16, color=BLUE, bold=True)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip(), size=13, color=BLUE, bold=True)
            i += 1
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[5:].strip(), size=12, color=DARK_BLUE, bold=True)
            i += 1
            continue
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            render_table(doc, parse_table(block))
            continue
        if re.match(r"^- ", line):
            while i < len(lines) and re.match(r"^- ", lines[i].rstrip()):
                p = doc.add_paragraph()
                attach_numbering(p, bullet_num)
                set_paragraph_spacing(p, after=8, line=1.167)
                add_inline(p, lines[i].rstrip()[2:].strip())
                i += 1
            continue
        if re.match(r"^\d+\. ", line):
            list_num = add_numbering(doc, bullet=False)
            while i < len(lines) and re.match(r"^\d+\. ", lines[i].rstrip()):
                p = doc.add_paragraph()
                attach_numbering(p, list_num)
                set_paragraph_spacing(p, after=8, line=1.167)
                add_inline(p, re.sub(r"^\d+\. ", "", lines[i].rstrip()))
                i += 1
            continue
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[2:-2], size=12, color=DARK_BLUE, bold=True)
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (
                not nxt
                or nxt.startswith(("## ", "### ", "#### ", "> ", "|", "- "))
                or re.match(r"^\d+\. ", nxt)
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=6, line=1.10)
        add_inline(p, " ".join(x.strip() for x in paragraph_lines))


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "해양 해커톤 모델 돌파구 딥리서치"
    props.subject = "P1/P2 코드 감사 및 두 사이클 연구 설계"
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "P1, P2, change point, F1, conditional FPCA, validation transport"

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    doc = Document()
    configure_styles(doc)
    setup_page(doc)
    set_core_properties(doc)
    bullet_num = add_numbering(doc, bullet=True)
    add_title_block(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    render_markdown(doc, lines, bullet_num)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
